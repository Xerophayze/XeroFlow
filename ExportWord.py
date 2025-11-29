# ExportWord.py

import sys
import subprocess
try:
    from importlib.metadata import distribution, PackageNotFoundError
except ImportError:  # Python <3.8 fallback
    from importlib_metadata import distribution, PackageNotFoundError
import logging

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s [%(levelname)s] %(message)s')

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import requests
from io import BytesIO
from tkinter import filedialog, Tk, messagebox
import os
import tempfile

# Import for Mermaid diagram support
try:
    import mermaid as md
    from mermaid.graph import Graph
    MERMAID_AVAILABLE = True
except ImportError:
    logging.warning("mermaid-py package not available. Mermaid diagrams will be rendered as code blocks.")
    MERMAID_AVAILABLE = False

### INLINE FORMATTING FUNCTIONS

def parse_inline_formatting(paragraph, text):
    """
    Parses inline GitHub-Flavored Markdown (GFM) formatting in the text and adds runs with appropriate styles.
    Supported formats include:
      - Combined formatting with strikethrough:
          • Bold with strikethrough: **~~text~~** or ~~**text**~~
          • Italic with strikethrough: *~~text~~* or ~~*text*~~
          • Bold, italic, and strikethrough: ***~~text~~*** or ~~***text***~~
      - Standard formatting:
          • Bold italic: ***text*** or ___text___, etc.
          • Bold: **text** or __text__
          • Italic: *text* or _text_
          • Underline via <u>text</u> (and combined with emphasis if wrapped accordingly)
          • Highlight via ==text== (and combined with other formatting)
          • Inline code, images, and links (with optional emphasis when wrapped)
    Note:
      • Plain ~~text~~ is processed as strikethrough only.
      • To have links styled as bold/italic, wrap the link markdown (e.g. **[link](url)**).
    """
    pattern = re.compile(
        # Combined formatting with strikethrough:
        r'(\*\*\*~~(.+?)~~\*\*\*)|'      # ***~~text~~***  => bold, italic, strikethrough
        r'(~~\*\*\*(.+?)\*\*\*~~)|'      # ~~***text***~~  => bold, italic, strikethrough
        r'(\*\*~~(.+?)~~\*\*)|'          # **~~text~~**      => bold, strikethrough
        r'(~~\*\*(.+?)\*\*~~)|'          # ~~**text**~~      => bold, strikethrough
        r'(\*~~(.+?)~~\*)|'              # *~~text~~*        => italic, strikethrough
        r'(~~\*(.+?)\*~~)|'              # ~~*text*~~        => italic, strikethrough
        
        # Underline combinations with emphasis:
        r'(<u>\*\*(.+?)\*\*</u>)|'       # <u>**text**</u>   => underline, bold
        r'(\*\*<u>(.+?)</u>\*\*)|'       # **<u>text</u>**   => underline, bold
        r'(<u>\*(.+?)\*</u>)|'           # <u>*text*</u>     => underline, italic
        r'(\*<u>(.+?)</u>\*)|'           # *<u>text</u>*     => underline, italic
        r'(<u>\*\*\*(.+?)\*\*\*</u>)|'   # <u>***text***</u> => underline, bold & italic
        r'(\*\*\*<u>(.+?)</u>\*\*\*)|'   # ***<u>text</u>*** => underline, bold & italic
        
        # Links with emphasis:
        r'(\*\*\[[^\]]+\]\([^\)]+\)\*\*)|'   # **[text](url)**   => bold link
        r'(\*\[[^\]]+\]\([^\)]+\)\*)|'       # *[text](url)*       => italic link
        r'(\*\*\*\[[^\]]+\]\([^\)]+\)\*\*\*)|'   # ***[text](url)*** => bold, italic link
        
        # Images with emphasis:
        r'(\*\*!\[[^\]]*\]\([^\)]+\)\*\*)|'     # **![alt](url)**   => bold image
        r'(\*!\[[^\]]*\]\([^\)]+\)\*)|'         # *![alt](url)*       => italic image
        r'(\*\*\*!!\[[^\]]*\]\([^\)]+\)\*\*\*)|' # ***!![alt](url)*** => bold, italic image
        
        # Original formatting patterns:
        r'(\*\*\*(.+?)\*\*\*)|'      # ***bold italic***
        r'(___(.+?)___)|'            # ___bold italic___
        r'(\*\*_(.+?)_\*\*)|'        # **_bold italic_**
        r'(__\*(.+?)\*__)|'          # __*bold italic*__
        r'(\*\*`(.+?)`\*\*)|'        # **`bold code`**
        r'(__`(.+?)`__)|'            # __`bold code`__
        r'(_`(.+?)`_)|'              # _`italic code`_
        r'(\*`(.+?)`\*)|'            # *`italic code`*
        r'(\*\*(.+?)\*\*)|'          # **bold**
        r'(__(.+?)__)|'              # __bold__
        r'(<u>(.+?)</u>)|'           # <u>underline</u>
        r'(==(.+?)==)|'              # ==highlight== (new pattern for highlighted text)
        r'(\*(.+?)\*)|'              # *italic*
        r'(_(.+?)_)|'                # _italic_
        r'(~~(.+?)~~)|'              # ~~strikethrough~~ (plain tilde: strike only)
        r'(`(.+?)`)|'                # `inline code`
        r'(!\[[^\]]*\]\([^\)]+\))|'  # image
        r'(\[[^\]]+\]\([^\)]+\))'    # link
    )
    last_index = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_index:
            paragraph.add_run(text[last_index:start])
        matched_text = match.group()

        # Combined formatting with strikethrough cases:
        if matched_text.startswith('***~~') or matched_text.startswith('~~***'):
            clean_text = matched_text[5:-5]
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.italic = True
            run.font.strike = True

        elif matched_text.startswith('**~~') or matched_text.startswith('~~**'):
            clean_text = matched_text[4:-4]
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.font.strike = True

        elif matched_text.startswith('*~~') or matched_text.startswith('~~*'):
            clean_text = matched_text[3:-3]
            run = paragraph.add_run(clean_text)
            run.italic = True
            run.font.strike = True

        # Underline combined with emphasis cases:
        elif matched_text.startswith('<u>***') or matched_text.startswith('***<u>'):
            if matched_text.startswith('<u>***'):
                clean_text = matched_text[6:-7]
            else:
                clean_text = re.sub(r'</?u>', '', matched_text[3:-3])
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.italic = True
            run.underline = True

        elif matched_text.startswith('<u>**') or matched_text.startswith('**<u>'):
            if matched_text.startswith('<u>**'):
                clean_text = matched_text[5:-6]
            else:
                clean_text = re.sub(r'</?u>', '', matched_text[2:-2])
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.underline = True

        elif matched_text.startswith('<u>*') or matched_text.startswith('*<u>'):
            if matched_text.startswith('<u>*'):
                clean_text = matched_text[4:-5]
            else:
                clean_text = re.sub(r'</?u>', '', matched_text[1:-1])
            run = paragraph.add_run(clean_text)
            run.italic = True
            run.underline = True

        # Links with emphasis (when wrapped with asterisks):
        elif matched_text.startswith('***['):
            link_content = matched_text[3:-3]
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', link_content)
            if m:
                link_text, url = m.groups()
                add_hyperlink(paragraph, link_text, url, bold=True, italic=True)
            else:
                paragraph.add_run(matched_text)

        elif matched_text.startswith('**['):
            link_content = matched_text[2:-2]
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', link_content)
            if m:
                link_text, url = m.groups()
                add_hyperlink(paragraph, link_text, url, bold=True)
            else:
                paragraph.add_run(matched_text)

        elif matched_text.startswith('*['):
            link_content = matched_text[1:-1]
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', link_content)
            if m:
                link_text, url = m.groups()
                add_hyperlink(paragraph, link_text, url, italic=True)
            else:
                paragraph.add_run(matched_text)

        # Images with emphasis cases (similar logic as links)...
        elif matched_text.startswith('***!!['):
            image_content = matched_text[5:-3]
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', image_content)
            if m:
                alt_text, url = m.groups()
                try:
                    response = requests.get(url)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(2))
                        run.bold = True
                        run.italic = True
                    else:
                        run = paragraph.add_run(f"[Image not available: {alt_text}]")
                        run.italic = True
                except Exception as e:
                    run = paragraph.add_run(f"[Image not available: {alt_text}]")
                    run.italic = True
            else:
                paragraph.add_run(matched_text)

        elif matched_text.startswith('**!['):
            image_content = matched_text[2:-2]
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', image_content)
            if m:
                alt_text, url = m.groups()
                try:
                    response = requests.get(url)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(2))
                        run.bold = True
                    else:
                        run = paragraph.add_run(f"[Image not available: {alt_text}]")
                        run.bold = True
                except Exception as e:
                    run = paragraph.add_run(f"[Image not available: {alt_text}]")
                    run.bold = True
            else:
                paragraph.add_run(matched_text)

        elif matched_text.startswith('*!['):
            image_content = matched_text[1:-1]
            m = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', image_content)
            if m:
                alt_text, url = m.groups()
                try:
                    response = requests.get(url)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(2))
                        run.italic = True
                    else:
                        run = paragraph.add_run(f"[Image not available: {alt_text}]")
                        run.italic = True
                except Exception as e:
                    run = paragraph.add_run(f"[Image not available: {alt_text}]")
                    run.italic = True
            else:
                paragraph.add_run(matched_text)

        # Existing combined formatting cases:
        elif matched_text.startswith('**_`') or matched_text.startswith('__*`'):
            clean_text = matched_text[4:-4]
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.italic = True
            run.font.name = 'Consolas'
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'D3D3D3')
            rPr.append(shd)

        elif matched_text.startswith('***') or matched_text.startswith('___') or \
             matched_text.startswith('**_') or matched_text.startswith('__*'):
            clean_text = re.sub(r'[*_]', '', matched_text)
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.italic = True

        elif matched_text.startswith('**`') or matched_text.startswith('__`'):
            clean_text = matched_text[3:-3]
            run = paragraph.add_run(clean_text)
            run.bold = True
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            paragraph._p.get_or_add_pPr().append(shading_elm)
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for border_type in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '1')
                border.set(qn('w:color'), 'auto')
                pBdr.append(border)
            pPr.append(pBdr)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)

        elif matched_text.startswith('_`') or matched_text.startswith('*`'):
            clean_text = matched_text[2:-2]
            run = paragraph.add_run(clean_text)
            run.italic = True
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            paragraph._p.get_or_add_pPr().append(shading_elm)
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            for border_type in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '1')
                border.set(qn('w:color'), 'auto')
                pBdr.append(border)
            pPr.append(pBdr)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)

        elif matched_text.startswith('**') or matched_text.startswith('__'):
            clean_text = matched_text[2:-2]
            run = paragraph.add_run(clean_text)
            run.bold = True

        elif matched_text.startswith('<u>'):
            clean_text = re.sub(r'</?u>', '', matched_text)
            run = paragraph.add_run(clean_text)
            run.underline = True

        elif matched_text.startswith('==') and matched_text.endswith('=='):
            clean_text = matched_text[2:-2]
            run = paragraph.add_run(clean_text)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Set highlight color to yellow

        elif matched_text.startswith('*') or matched_text.startswith('_'):
            clean_text = matched_text[1:-1]
            run = paragraph.add_run(clean_text)
            run.italic = True

        # Plain tilde: only strikethrough.
        elif matched_text.startswith('~~'):
            clean_text = matched_text[2:-2]
            run = paragraph.add_run(clean_text)
            run.font.strike = True

        elif matched_text.startswith('`'):
            clean_text = matched_text[1:-1]
            run = paragraph.add_run(clean_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'E0E0E0')
            paragraph._p.get_or_add_pPr().append(shading_elm)

        # Image (plain)
        elif matched_text.startswith('!['):
            alt_text, url = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', matched_text).groups()
            try:
                response = requests.get(url)
                if response.status_code == 200:
                    image_stream = BytesIO(response.content)
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(2))
            except Exception as e:
                run = paragraph.add_run(f"[Image not available: {alt_text}]")
                run.italic = True

        # Link (plain)
        elif matched_text.startswith('['):
            link_text, url = re.match(r'\[([^\]]+)\]\(([^)]+)\)', matched_text).groups()
            add_hyperlink(paragraph, link_text, url)

        else:
            paragraph.add_run(matched_text)
        last_index = end
    if last_index < len(text):
        paragraph.add_run(text[last_index:])

def add_hyperlink(paragraph, text, url, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    if bold:
        b_elem = OxmlElement('w:b')
        b_elem.set(qn('w:val'), 'true')
        rPr.append(b_elem)
    if italic:
        i_elem = OxmlElement('w:i')
        i_elem.set(qn('w:val'), 'true')
        rPr.append(i_elem)
    new_run.append(rPr)
    w_t = OxmlElement('w:t')
    w_t.text = text
    new_run.append(w_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_heading_style(paragraph, level):
    sizes = {1: 20, 2: 18, 3: 16, 4: 14, 5: 12, 6: 11}
    font_size = sizes.get(level, 10)
    colors = {
        1: RGBColor(0, 0, 139),
        2: RGBColor(0, 0, 139),
        3: RGBColor(47, 47, 47),
        4: RGBColor(47, 47, 47),
        5: RGBColor(47, 47, 47),
        6: RGBColor(47, 47, 47)
    }
    paragraph.style = f'Heading {level}'
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
        run.font.name = 'Calibri'
        run.bold = True
    if level <= 2:
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8' if level == 1 else '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True

def add_code_block(document, code_text, language):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'E0E0E0')
    paragraph._p.get_or_add_pPr().append(shading_elm)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_type in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_type}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), 'auto')
        pBdr.append(border)
    pPr.append(pBdr)
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)

def add_bookmark(paragraph, bookmark_name, bookmark_id):
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    runs = paragraph.runs
    if runs:
        first_run = runs[0]
        first_run._r.insert(0, bookmark_start)
        last_run = runs[-1]
        last_run._r.append(bookmark_end)
    else:
        run = paragraph.add_run()
        run.text = ''
        run._r.insert(0, bookmark_start)
        run._r.append(bookmark_end)

def add_hyperlink(paragraph, text, url, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    if bold:
        b_elem = OxmlElement('w:b')
        b_elem.set(qn('w:val'), 'true')
        rPr.append(b_elem)
    if italic:
        i_elem = OxmlElement('w:i')
        i_elem.set(qn('w:val'), 'true')
        rPr.append(i_elem)
    new_run.append(rPr)
    w_t = OxmlElement('w:t')
    w_t.text = text
    new_run.append(w_t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def parse_table(lines, current_index):
    """
    Parse a markdown table and return the table data and the index after the table.
    """
    if current_index >= len(lines):
        return None, current_index

    # Check if we have a valid table header
    header_line = lines[current_index].strip()
    if not header_line.startswith('|'):
        return None, current_index

    # Extract headers
    headers = [cell.strip() for cell in header_line.strip('|').split('|')]
    current_index += 1

    # Check for separator line, but don't require it
    if current_index < len(lines):
        separator_line = lines[current_index].strip()
        if separator_line.startswith('|') and all('-' in cell for cell in separator_line.strip('|').split('|')):
            current_index += 1

    # Extract data rows
    data = []
    while current_index < len(lines):
        line = lines[current_index].strip()
        if not line.startswith('|'):
            break
        row = [cell.strip() for cell in line.strip('|').split('|')]
        if len(row) == len(headers):
            data.append(row)
        current_index += 1

    if not data:  # No data rows found
        return None, current_index

    return {'headers': headers, 'data': data}, current_index

def insert_table(document, table_data):
    """
    Add a table to the Word document.
    """
    if not table_data or 'headers' not in table_data or 'data' not in table_data:
        return

    # Create table
    table = document.add_table(rows=1, cols=len(table_data['headers']))
    table.style = 'Table Grid'

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(table_data['headers']):
        header_cells[i].text = ""
        paragraph = header_cells[i].paragraphs[0]
        parse_inline_formatting(paragraph, header)
        paragraph.runs[0].bold = True

    # Add data rows
    for row_data in table_data['data']:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            # Split content by <br> tags
            lines = cell_data.split('<br>')
            row_cells[i].text = ""  # Clear default paragraph
            
            # Process each line
            for line_idx, line in enumerate(lines):
                line = line.strip()
                if line_idx > 0:  # Add new paragraph for each line after the first
                    paragraph = row_cells[i].add_paragraph()
                else:
                    paragraph = row_cells[i].paragraphs[0]
                
                # Handle list items with proper indentation
                if line.strip().startswith(('-', '*', '+')):
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    parse_inline_formatting(paragraph, line)
                elif line.strip().startswith(tuple(f"{n}." for n in range(1, 10))):
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    parse_inline_formatting(paragraph, line)
                else:
                    parse_inline_formatting(paragraph, line)

    # Add space after table
    document.add_paragraph()

### END TABLE FUNCTIONS

### CHART FUNCTIONS

def process_chart_code(code_text):
    """
    Process chart code and return the image bytes.
    Args:
        code_text (str): The chart configuration code
    Returns:
        bytes: The chart image bytes
    """
    from quickchart import QuickChart
    import json
    import re
    try:
        logging.info("\n[DEBUG Chart] Processing chart code:")
        logging.info(f"Raw code text: {code_text}")
        
        # Pre-process the JSON to handle complex cases
        
        # 1. Handle Mermaid diagrams embedded in JSON
        mermaid_match = re.search(r'("mermaid"\s*:\s*")([\s\S]*?)("\s*\n*\s*})', code_text)
        if mermaid_match:
            mermaid_content = mermaid_match.group(2).strip()
            return process_mermaid_diagram(mermaid_content)
        
        # 2. Remove all JavaScript functions (callbacks, formatters, etc.)
        # Remove callback functions
        code_text = re.sub(r'"callback":\s*"function\(.*?\)\s*{[^}]*}"', '', code_text, flags=re.DOTALL)
        # Remove formatter functions
        code_text = re.sub(r'"formatter":\s*\(.*?\)\s*=>\s*{[^}]*},?', '', code_text, flags=re.DOTALL)
        # Remove tooltip callbacks
        code_text = re.sub(r'"tooltips":\s*{[^}]*"callbacks":\s*{[^}]*}}', '', code_text, flags=re.DOTALL)
        # Remove annotation callbacks
        code_text = re.sub(r'"annotation":\s*{[^}]*"annotations":\s*\[[^]]*\]}', '', code_text, flags=re.DOTALL)
        
        # 3. Remove trailing commas
        code_text = re.sub(r',(\s*})', r'\1', code_text)
        
        # 4. Handle duplicated keys
        code_text = re.sub(r'("xAxes": \[[\s\S]*?"scaleLabel": {[\s\S]*?})\s*"xAxes":', r'\1,', code_text, flags=re.DOTALL)
        
        logging.info(f"Preprocessed code text: {code_text}")
        
        # Parse the chart configuration
        config = json.loads(code_text)
        logging.info(f"Parsed config: {json.dumps(config, indent=2)}")
        
        # Create QuickChart instance
        qc = QuickChart()
        qc.width = config.get('width', 500)
        qc.height = config.get('height', 300)
        qc.config = config.get('config', {})
        
        logging.info(f"QuickChart config - Width: {qc.width}, Height: {qc.height}")
        logging.info(f"Chart config: {json.dumps(qc.config, indent=2)}")
        
        # Prepare the request data
        post_data = {
            'chart': qc.config,
            'width': qc.width,
            'height': qc.height,
            'format': 'png'
        }
        
        # Make the request with proper headers
        headers = {
            'Content-Type': 'application/json',
            'Accept': 'image/png'
        }
        
        response = requests.post(
            'https://quickchart.io/chart', 
            json=post_data, 
            headers=headers,
            timeout=30
        )
        
        # Check if response is an image (even with error status code)
        content_type = response.headers.get('Content-Type', '')
        if content_type.startswith('image/'):
            logging.info(f"Successfully generated chart image of size: {len(response.content)} bytes")
            return response.content
        else:
            logging.error(f"Unexpected response type: {content_type}")
            logging.error(f"Response body: {response.text}")
            return None
            
    except json.JSONDecodeError as je:
        logging.error(f"JSON parsing error: {je}")
        logging.error(f"Invalid JSON: {code_text}")
        return None
    except Exception as e:
        logging.error(f"Error processing chart code: {e}")
        logging.error(f"Stack trace:", exc_info=True)
        return None

def process_mermaid_diagram(mermaid_code):
    """
    Process Mermaid diagram code and return the image bytes.
    Args:
        mermaid_code (str): The Mermaid diagram code
    Returns:
        bytes: The diagram image bytes or None if processing fails
    """
    if not MERMAID_AVAILABLE:
        logging.warning("Mermaid package not available. Cannot render diagram.")
        return None
    
    try:
        logging.info("\n[DEBUG Mermaid] Processing Mermaid diagram:")
        logging.info(f"Raw code text: {mermaid_code}")

        # Pre-process to quote subgraph titles and node labels that are not already quoted.
        # This prevents parse errors for titles/labels with special characters like (), /, etc.
        import re

        # Robust Mermaid Pre-processing Pipeline

        # 1. Fix misplaced <br> tags. e.g., `Node(text)<br>more` becomes `Node(text\nmore)`
        # This regex finds a node definition, captures the part inside brackets, and any trailing <br> tags.
        # mermaid_code = re.sub(r'([(\[{][^)\]}]*[\])}])(<br\s*/?>)+(.*)', lambda m: f"{m.group(1)[:-1]}\n{m.group(3)}{m.group(1)[-1]}", mermaid_code, flags=re.IGNORECASE)

        # 2. Define the main processing function for node labels.
        # def process_node_label(match):
        #     node_id, opener, label, closer = match.groups()

        #     # 2a. Convert all <br> tags to newline characters.
        #     label = re.sub(r'<br\s*/?>', '\n', label, flags=re.IGNORECASE)

        #     # 2b. Escape any existing quotes inside the label.
        #     label = label.replace('"', '\\"')

        #     # 2c. Return the node with the label now safely quoted.
        #     return f'{node_id}{opener}"{label}"{closer}'

        # # 3. Apply the label processing to all nodes.
        # mermaid_code = re.sub(r'([a-zA-Z0-9_]+)([(\[{])(.*?)([\])}])', process_node_label, mermaid_code)

        # # 4. Define a processing function for subgraph titles.
        # def process_subgraph_title(match):
        #     keyword, title = match.groups()
        #     title = title.strip()
        #     # Quote if not already quoted
        #     if not (title.startswith('"') and title.endswith('"')):
        #         title = title.replace('"', '\\"') # Escape internal quotes
        #         return f'{keyword}"{title}"'
        #     return match.group(0) # Return original if already quoted

        # --- Start of Mermaid Pre-processing ---

        def preprocess_mermaid(code):
            # This simplified function trusts the AI to generate mostly correct syntax based on
            # the improved examples. It only performs the most critical and safe corrections.

            # 1. Fix subgraph styles defined by title instead of ID.
            # This is a common AI error and is safe to fix programmatically.
            subgraph_titles = re.findall(r'subgraph "(.*?)"', code)
            for title in subgraph_titles:
                subgraph_id = re.sub(r'[^\w]+', '_', title)
                # Update subgraph declaration to use an ID, e.g., subgraph My_ID ["My Title"]
                code = code.replace(f'subgraph "{title}"', f'subgraph {subgraph_id} ["{title}"]')
                # Update the corresponding style definition to use the new ID
                code = code.replace(f'style "{title}"', f'style {subgraph_id}')

            # 2. Standardize link syntax: Replace --- with -- for labels.
            code = re.sub(r'--- "(.*?)" ---', r'-- "\1" --', code)

            return code

        mermaid_code = preprocess_mermaid(mermaid_code)

        # --- End of Mermaid Pre-processing ---

        # Use Kroki SVG as primary method (no width limits), fall back to mermaid.ink if needed
        try:
            # Try Kroki SVG first - it's vector-based with no width limitations
            # Default to local Kroki instance on searxng.local, fall back to public kroki.io if needed
            kroki_base = os.environ.get("KROKI_SERVER", "http://searxng.local:8000").rstrip('/')
            kroki_svg_url = f"{kroki_base}/mermaid/svg"
            headers = {
                'Accept': 'image/svg+xml',
                'Content-Type': 'text/plain'
            }
            logging.info(f"Attempting Kroki SVG rendering at: {kroki_svg_url}")
            
            # Try Kroki with timeout handling (reduced timeout for faster fallback)
            try:
                response = requests.post(kroki_svg_url, data=mermaid_code.encode('utf-8'), headers=headers, timeout=10)
            except requests.exceptions.Timeout:
                logging.warning("Kroki request timed out after 10 seconds. Falling back to mermaid.ink...")
                response = None
            except requests.exceptions.RequestException as req_err:
                logging.warning(f"Kroki request failed: {req_err}. Falling back to mermaid.ink...")
                response = None
            
            if response and response.status_code == 200 and response.headers.get('Content-Type', '').startswith('image/svg'):
                logging.info(f"Successfully generated Mermaid SVG via Kroki. SVG size: {len(response.content)} bytes")
                
                # Convert SVG to high-resolution PNG using Kroki's PNG endpoint
                # This avoids SVG compatibility issues with python-docx and text rendering problems
                try:
                    # Request PNG version from Kroki at high resolution
                    kroki_png_url = f"{kroki_base}/mermaid/png"
                    png_headers = {
                        'Accept': 'image/png',
                        'Content-Type': 'text/plain'
                    }
                    logging.info("Converting to high-resolution PNG via Kroki...")
                    
                    # Try PNG conversion with timeout handling (reduced timeout for faster fallback)
                    try:
                        png_response = requests.post(kroki_png_url, data=mermaid_code.encode('utf-8'), headers=png_headers, timeout=10)
                    except (requests.exceptions.Timeout, requests.exceptions.RequestException) as png_err:
                        logging.warning(f"Kroki PNG conversion timed out or failed after 10 seconds: {png_err}")
                        png_response = None
                    
                    if png_response and png_response.status_code == 200 and png_response.headers.get('Content-Type', '').startswith('image/'):
                        from PIL import Image
                        import io
                        
                        # Open PNG and get dimensions
                        png_image = Image.open(BytesIO(png_response.content))
                        logging.info(f"Kroki PNG size: {png_image.size[0]}x{png_image.size[1]} pixels, {len(png_response.content):,} bytes")
                        
                        # Convert to RGB for JPEG
                        if png_image.mode in ('RGBA', 'LA', 'P'):
                            rgb_image = Image.new('RGB', png_image.size, (255, 255, 255))
                            if png_image.mode == 'P':
                                png_image = png_image.convert('RGBA')
                            rgb_image.paste(png_image, mask=png_image.split()[-1] if png_image.mode in ('RGBA', 'LA') else None)
                            png_image = rgb_image
                        elif png_image.mode != 'RGB':
                            png_image = png_image.convert('RGB')
                        
                        # Save as JPEG with 80% quality
                        jpg_buffer = io.BytesIO()
                        png_image.save(jpg_buffer, format='JPEG', quality=80, optimize=True, progressive=True)
                        jpg_bytes = jpg_buffer.getvalue()
                        
                        logging.info(f"Converted to JPEG: {len(jpg_bytes):,} bytes (quality=80)")
                        return jpg_bytes
                    else:
                        status = png_response.status_code if png_response else "timeout/error"
                        logging.warning(f"Kroki PNG conversion failed: {status}")
                        # Fall through to mermaid.ink fallback
                except Exception as e:
                    logging.error(f"Error converting via Kroki PNG: {e}")
                    # Fall through to mermaid.ink fallback
            else:
                status = response.status_code if response else "timeout/error"
                logging.warning(f"Kroki SVG failed. Status: {status}. Falling back to mermaid.ink...")
            
            # Fallback to mermaid.ink - use SVG endpoint for better quality
            import base64
            graphbytes = mermaid_code.encode("utf8")
            base64_bytes = base64.urlsafe_b64encode(graphbytes)
            base64_string = base64_bytes.decode("ascii")
            svg_url = f"https://mermaid.ink/svg/{base64_string}"
            
            logging.info("Attempting mermaid.ink fallback (SVG)...")
            response = requests.get(svg_url, timeout=15)
            if response.status_code == 200 and len(response.content) > 100:
                logging.info(f"Successfully generated Mermaid SVG using mermaid.ink. SVG size: {len(response.content)} bytes")
                
                # Convert SVG to high-resolution PNG, then to JPEG
                try:
                    from PIL import Image
                    import io
                    
                    # Try multiple SVG conversion methods
                    png_data = None
                    conversion_method = None
                    
                    # Method 1: Try cairosvg (best quality)
                    try:
                        import cairosvg
                        png_data = cairosvg.svg2png(
                            bytestring=response.content,
                            scale=4.0  # 4x resolution for high quality
                        )
                        conversion_method = "cairosvg (4x scale)"
                    except ImportError:
                        logging.info("cairosvg not available, trying alternative methods...")
                    except Exception as e:
                        logging.warning(f"cairosvg conversion failed: {e}")
                    
                    # Method 2: Try svglib + reportlab
                    if not png_data:
                        try:
                            from svglib.svglib import svg2rlg
                            from reportlab.graphics import renderPM
                            import tempfile
                            
                            # Save SVG to temp file
                            with tempfile.NamedTemporaryFile(suffix='.svg', delete=False) as tmp_svg:
                                tmp_svg.write(response.content)
                                tmp_svg_path = tmp_svg.name
                            
                            # Convert SVG to ReportLab drawing
                            drawing = svg2rlg(tmp_svg_path)
                            if drawing:
                                # Scale up for higher resolution
                                drawing.width *= 4
                                drawing.height *= 4
                                drawing.scale(4, 4)
                                
                                # Render to PNG
                                png_data = renderPM.drawToString(drawing, fmt='PNG')
                                conversion_method = "svglib (4x scale)"
                            
                            # Clean up temp file
                            os.unlink(tmp_svg_path)
                        except ImportError:
                            logging.info("svglib not available, trying alternative methods...")
                        except Exception as e:
                            logging.warning(f"svglib conversion failed: {e}")
                    
                    # Method 3: Use wand (ImageMagick) if available
                    if not png_data:
                        try:
                            from wand.image import Image as WandImage
                            with WandImage(blob=response.content, format='svg', resolution=300) as img:
                                img.format = 'png'
                                png_data = img.make_blob()
                            conversion_method = "wand (300 DPI)"
                        except ImportError:
                            logging.info("wand not available")
                        except Exception as e:
                            logging.warning(f"wand conversion failed: {e}")
                    
                    if png_data:
                        # Open the PNG
                        png_image = Image.open(BytesIO(png_data))
                        logging.info(f"Converted SVG to PNG using {conversion_method}: {png_image.size[0]}x{png_image.size[1]} pixels")
                        
                        # Convert to RGB (JPG doesn't support transparency)
                        if png_image.mode in ('RGBA', 'LA', 'P'):
                            rgb_image = Image.new('RGB', png_image.size, (255, 255, 255))
                            if png_image.mode == 'P':
                                png_image = png_image.convert('RGBA')
                            rgb_image.paste(png_image, mask=png_image.split()[-1] if png_image.mode in ('RGBA', 'LA') else None)
                            png_image = rgb_image
                        elif png_image.mode != 'RGB':
                            png_image = png_image.convert('RGB')
                        
                        # Save as JPEG with 80% quality (same as Kroki path)
                        jpg_buffer = io.BytesIO()
                        png_image.save(jpg_buffer, format='JPEG', quality=80, optimize=True, progressive=True)
                        jpg_bytes = jpg_buffer.getvalue()
                        
                        logging.info(f"Converted to JPEG: {len(jpg_bytes):,} bytes (quality=80)")
                        return jpg_bytes
                    else:
                        # No SVG converter available, fall back to low-res PNG
                        logging.warning("No SVG conversion library available (cairosvg, svglib, or wand). Falling back to low-resolution PNG...")
                        
                except Exception as conv_error:
                    logging.error(f"Error converting SVG to PNG/JPEG: {conv_error}")
                
                # Fallback: Modify SVG dimensions and try to render at higher resolution
                try:
                    from PIL import Image
                    import io
                    import re
                    
                    # Modify the SVG to render at higher resolution
                    svg_text = response.content.decode('utf-8')
                    
                    # Extract viewBox dimensions
                    viewbox_match = re.search(r'viewBox="([^"]+)"', svg_text)
                    if viewbox_match:
                        viewbox_parts = viewbox_match.group(1).split()
                        if len(viewbox_parts) == 4:
                            vb_width = float(viewbox_parts[2])
                            vb_height = float(viewbox_parts[3])
                            
                            # Set explicit large dimensions (6x scale for better quality)
                            new_width = int(vb_width * 6)
                            new_height = int(vb_height * 6)
                            
                            # Remove existing width/height and add new ones
                            svg_text = re.sub(r'\s*width="[^"]*"', '', svg_text)
                            svg_text = re.sub(r'\s*height="[^"]*"', '', svg_text)
                            svg_text = re.sub(r'<svg', f'<svg width="{new_width}" height="{new_height}"', svg_text, count=1)
                            
                            logging.info(f"Modified SVG dimensions to {new_width}x{new_height} (6x scale)")
                            
                            # Try to convert the modified SVG with cairosvg if available
                            try:
                                import cairosvg
                                png_data = cairosvg.svg2png(bytestring=svg_text.encode('utf-8'))
                                
                                # Open and convert to JPEG
                                png_image = Image.open(BytesIO(png_data))
                                logging.info(f"Rendered SVG to PNG: {png_image.size[0]}x{png_image.size[1]} pixels")
                                
                                # Convert to RGB
                                if png_image.mode in ('RGBA', 'LA', 'P'):
                                    rgb_image = Image.new('RGB', png_image.size, (255, 255, 255))
                                    if png_image.mode == 'P':
                                        png_image = png_image.convert('RGBA')
                                    rgb_image.paste(png_image, mask=png_image.split()[-1] if png_image.mode in ('RGBA', 'LA') else None)
                                    png_image = rgb_image
                                elif png_image.mode != 'RGB':
                                    png_image = png_image.convert('RGB')
                                
                                # Save as JPEG
                                jpg_buffer = io.BytesIO()
                                png_image.save(jpg_buffer, format='JPEG', quality=85, optimize=True, progressive=True)
                                jpg_bytes = jpg_buffer.getvalue()
                                
                                logging.info(f"Converted to high-resolution JPEG: {len(jpg_bytes):,} bytes (quality=85)")
                                return jpg_bytes
                            except Exception as cairo_error:
                                logging.warning(f"Cairo conversion failed: {cairo_error}. Falling back to PNG upscaling...")
                    
                    # Final fallback: Download low-res PNG and upscale
                    png_url = f"https://mermaid.ink/img/{base64_string}"
                    png_response = requests.get(png_url, timeout=15)
                    if png_response.status_code == 200:
                        logging.info(f"Downloaded low-resolution PNG: {len(png_response.content)} bytes")
                        
                        # Open the low-res PNG
                        low_res_image = Image.open(BytesIO(png_response.content))
                        original_size = low_res_image.size
                        logging.info(f"Original size: {original_size[0]}x{original_size[1]} pixels")
                        
                        # Upscale 6x using high-quality resampling for better quality
                        scale_factor = 6
                        new_size = (original_size[0] * scale_factor, original_size[1] * scale_factor)
                        high_res_image = low_res_image.resize(new_size, Image.Resampling.LANCZOS)
                        logging.info(f"Upscaled to: {new_size[0]}x{new_size[1]} pixels ({scale_factor}x)")
                        
                        # Convert to RGB
                        if high_res_image.mode in ('RGBA', 'LA', 'P'):
                            rgb_image = Image.new('RGB', high_res_image.size, (255, 255, 255))
                            if high_res_image.mode == 'P':
                                high_res_image = high_res_image.convert('RGBA')
                            rgb_image.paste(high_res_image, mask=high_res_image.split()[-1] if high_res_image.mode in ('RGBA', 'LA') else None)
                            high_res_image = rgb_image
                        elif high_res_image.mode != 'RGB':
                            high_res_image = high_res_image.convert('RGB')
                        
                        # Save as JPEG with higher quality
                        jpg_buffer = io.BytesIO()
                        high_res_image.save(jpg_buffer, format='JPEG', quality=85, optimize=True, progressive=True)
                        jpg_bytes = jpg_buffer.getvalue()
                        
                        logging.info(f"Converted to high-resolution JPEG: {len(jpg_bytes):,} bytes (quality=85, {scale_factor}x upscale)")
                        return jpg_bytes
                    else:
                        logging.error("Low-resolution PNG fallback also failed")
                        return None
                except Exception as fallback_error:
                    logging.error(f"Fallback processing failed: {fallback_error}")
                    return None
            else:
                logging.error(f"mermaid.ink also failed. Status: {response.status_code}")
                return None
                
        except Exception as e:
            logging.error(f"Error processing Mermaid diagram: {e}")
            return None
    except Exception as e:
        logging.error(f"Error processing Mermaid diagram: {e}")
        logging.error(f"Stack trace:", exc_info=True)
        return None

def add_code_block(document, code_text, language=None):
    """
    Add a code block to the Word document with syntax highlighting.
    Args:
        document: The Word document
        code_text: The code content
        language: The programming language (for reference)
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    # Add a light gray background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    paragraph._element.get_or_add_pPr().append(shading_elm)

def insert_chart_image(document, image_bytes):
    """
    Insert a chart image into the Word document.
    Args:
        document: The Word document
        image_bytes: The image bytes to insert
    Returns:
        paragraph: The paragraph containing the image
    """
    try:
        paragraph = document.add_paragraph()
        if image_bytes:
            # Detect image format from bytes to use correct file extension
            if image_bytes.startswith(b'<svg') or image_bytes.startswith(b'<?xml'):
                suffix = '.svg'
            elif image_bytes.startswith(b'\x89PNG'):
                suffix = '.png'
            elif image_bytes.startswith(b'\xff\xd8\xff'):
                suffix = '.jpg'
            else:
                suffix = '.png'  # Default fallback
            
            # Save image bytes to a temporary file with correct extension
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp_file:
                tmp_file.write(image_bytes)
                temp_path = tmp_file.name
            
            try:
                # Add picture from the temporary file
                paragraph.add_run().add_picture(temp_path, width=Inches(6))
                # Clean up the temporary file
                os.unlink(temp_path)
            except Exception as e:
                logging.error(f"Error adding picture from temp file: {e}")
                os.unlink(temp_path)  # Clean up even if there's an error
                raise
        return paragraph
    except Exception as e:
        try:
            import importlib.metadata
            if importlib.metadata.version('mutagen') is None:
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'mutagen'])
        except Exception as e:
            logging.error(f"Error installing mutagen: {e}")
        return document.add_paragraph("[Error: Failed to insert chart]")

def save_image_to_subdirectory(image_bytes, output_path, image_type, image_index):
    """
    Save an image to the ChartsAndDiagrams subdirectory.
    Args:
        image_bytes: The image bytes to save
        output_path: The path where the main document will be saved
        image_type: Type of image ('chart' or 'mermaid')
        image_index: Index number for the image
    Returns:
        str: The path where the image was saved, or None if failed
    """
    try:
        if not image_bytes or not output_path:
            return None
            
        # Get the directory where the document will be saved
        doc_dir = os.path.dirname(output_path)
        
        # Create ChartsAndDiagrams subdirectory
        charts_dir = os.path.join(doc_dir, "ChartsAndDiagrams")
        os.makedirs(charts_dir, exist_ok=True)
        
        # Generate filename based on document name and image index
        doc_name = os.path.splitext(os.path.basename(output_path))[0]
        
        # Detect file type from image bytes
        if image_bytes.startswith(b'<svg') or image_bytes.startswith(b'<?xml'):
            file_ext = 'svg'
        elif image_bytes.startswith(b'\x89PNG'):
            file_ext = 'png'
        elif image_bytes.startswith(b'\xff\xd8\xff'):
            file_ext = 'jpg'
        else:
            # Default based on image type
            file_ext = 'svg' if image_type == 'mermaid' else 'png'
            
        image_filename = f"{doc_name}_{image_type}_{image_index:03d}.{file_ext}"
        image_path = os.path.join(charts_dir, image_filename)
        
        # Save the image
        with open(image_path, 'wb') as f:
            f.write(image_bytes)
            
        logging.info(f"[DEBUG ExportWord] Saved {image_type} image to: {image_path}")
        return image_path
        
    except Exception as e:
        logging.error(f"[DEBUG ExportWord] Failed to save {image_type} image: {e}")
        return None

### END CHART FUNCTIONS

### LIST FUNCTIONS

def get_list_level(line):
    """
    Calculate the indentation level of a list item.
    Args:
        line (str): The line containing the list item
    Returns:
        int: The indentation level (0-based)
    """
    indent = 0
    for char in line:
        if char == ' ':
            indent += 1
        elif char == '\t':
            indent += 2
        else:
            break
    return indent // 2

def handle_list_item(document, line, list_level=0, list_counters=None):
    """
    Process a list item line and create a properly formatted paragraph.
    Args:
        document: The Word document
        line: The line containing the list item
        list_level: The indentation level of the list item
        list_counters: Dictionary tracking the current number for each list level
    """
    # Check for numbered list item (1., 2., etc.)
    numbered_match = re.match(r'^(\s*)(\d+\.)\s+(.*)$', line)
    if numbered_match:
        number, content = numbered_match.groups()[1:]
        
        # Create a regular paragraph instead of using List Number style
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25 * list_level)
        
        # If we're manually tracking list numbers, use the tracked number
        # Otherwise, use the number from the markdown
        if list_counters is not None:
            # If this is a new level, initialize the counter
            if list_level not in list_counters:
                list_counters[list_level] = 1
            
            # Add the number as text
            run = paragraph.add_run(f"{list_counters[list_level]}. ")
            
            # Increment the counter for this level
            list_counters[list_level] += 1
            
            # Reset all deeper level counters when we encounter a higher level item
            deeper_levels = [level for level in list_counters.keys() if level > list_level]
            for level in deeper_levels:
                list_counters[level] = 1
        else:
            # Just use the number from the markdown
            run = paragraph.add_run(number + " ")
        
        parse_inline_formatting(paragraph, content)
        return paragraph
    
    # Check for alphabetic subitem (a., b., etc.)
    alpha_match = re.match(r'^(>+)\s*([a-z]\.)\s+(.*)$', line)
    if alpha_match:
        quote_level = len(alpha_match.group(1))
        letter, content = alpha_match.groups()[1:]
        paragraph = document.add_paragraph()
        # Set indentation for subitems
        paragraph.paragraph_format.left_indent = Inches(0.25 * list_level + 0.25)
        # Add the letter as text
        run = paragraph.add_run(f"{letter} ")
        parse_inline_formatting(paragraph, content)
        return paragraph
    
    # Check for bullet list item (-, *, +)
    bullet_match = re.match(r'^(\s*)([-*+])\s+(.*)$', line)
    if bullet_match:
        content = bullet_match.groups()[2]
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.paragraph_format.left_indent = Inches(0.25 * list_level)
        parse_inline_formatting(paragraph, content)
        return paragraph
    
    # Default case - just a regular paragraph
    paragraph = document.add_paragraph()
    parse_inline_formatting(paragraph, line)
    return paragraph

### CONVERSION FUNCTION

def convert_markdown_to_docx(markdown_text, output_path=None, formatting_enabled=False):
    """
    Converts GitHub-Flavored Markdown (GFM) text to a Word document.
    If output_path is provided, saves to that location.
    Otherwise, opens a file dialog to choose the save location.
    
    Args:
        markdown_text (str): The markdown text to convert
        output_path (str, optional): Path where to save the Word document. If None, shows file dialog.
        formatting_enabled (bool, optional): Whether to apply markdown formatting. If False, saves as plain text.
    """
    logging.info("\n[DEBUG ExportWord] Starting conversion:")
    logging.info("-" * 80)
    logging.info(f"Formatting enabled: {formatting_enabled}")
    logging.info(f"Output path: {output_path}")
    logging.info(f"Content type: {type(markdown_text)}")
    logging.info(f"Content length: {len(markdown_text)}")
    logging.info("First 500 chars of content:")
    logging.info(markdown_text[:500])
    logging.info("-" * 80)

    file_path = output_path
    if file_path is None:
        root = Tk()
        root.withdraw()
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            title="Save as"
        )
        if not file_path:
            messagebox.showwarning("No Save Location", "No file was selected. Operation cancelled.")
            return
    try:
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

        in_code_block = False
        code_block_language = ''
        code_block_content = []
        
        # Dictionary to track list numbering for each level
        list_counters = {}
        
        # Counters for chart and diagram images
        chart_counter = 0
        mermaid_counter = 0

        if not formatting_enabled:
            logging.info("[DEBUG ExportWord] Formatting disabled, adding text as-is")
            paragraph = document.add_paragraph()
            paragraph.add_run(markdown_text)
        else:
            logging.info("[DEBUG ExportWord] Formatting enabled, processing markdown")
            lines = markdown_text.split('\n')
            logging.info(f"[DEBUG ExportWord] Number of lines to process: {len(lines)}")
            bookmark_id = 0
            index = 0
            while index < len(lines):
                line = lines[index]
                stripped_line = line.strip()
                if in_code_block:
                    if stripped_line.startswith('```'):
                        code_text = '\n'.join(code_block_content)
                        logging.info(f"[DEBUG ExportWord] Processing code block with language: '{code_block_language}'")
                        if code_block_language == 'chart':
                            logging.info("[DEBUG ExportWord] Processing chart code block")
                            image_bytes = process_chart_code(code_text)
                            if image_bytes:
                                logging.info("[DEBUG ExportWord] Successfully generated chart image")
                                chart_counter += 1
                                # Save image to subdirectory
                                save_image_to_subdirectory(image_bytes, file_path, 'chart', chart_counter)
                                # Insert image into document
                                insert_chart_image(document, image_bytes)
                            else:
                                logging.error("[DEBUG ExportWord] Failed to generate chart image, falling back to code block")
                                add_code_block(document, code_text, code_block_language)
                        elif code_block_language == 'mermaid':
                            logging.info("[DEBUG ExportWord] Processing mermaid code block")
                            image_bytes = process_mermaid_diagram(code_text)
                            if image_bytes:
                                logging.info("[DEBUG ExportWord] Successfully generated mermaid diagram image")
                                mermaid_counter += 1
                                # Save image to subdirectory
                                save_image_to_subdirectory(image_bytes, file_path, 'mermaid', mermaid_counter)
                                # Insert image into document
                                insert_chart_image(document, image_bytes)
                            else:
                                logging.error("[DEBUG ExportWord] Failed to generate mermaid diagram, falling back to code block")
                                add_code_block(document, code_text, code_block_language)
                        else:
                            add_code_block(document, code_text, code_block_language)
                        in_code_block = False
                        code_block_language = ''
                        code_block_content = []
                    else:
                        code_block_content.append(line)
                    index += 1
                    continue
                if stripped_line.startswith('```'):
                    in_code_block = True
                    code_block_language = stripped_line[3:].strip().lower()
                    logging.info(f"[DEBUG ExportWord] Code block started with language: '{code_block_language}'")
                    index += 1
                    continue
                # Check for page break tag
                if stripped_line == '<pbreak>':
                    logging.info("[DEBUG ExportWord] Adding page break")
                    add_page_break(document)
                    index += 1
                    continue
                if re.match(r'^(\*\*\*|---|___)$', stripped_line):
                    paragraph = document.add_paragraph()
                    run = paragraph.add_run()
                    pPr = paragraph._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), 'auto')
                    pBdr.append(bottom)
                    pPr.append(pBdr)
                    index += 1
                    continue
                if stripped_line.startswith('|') and stripped_line.endswith('|'):
                    parsed = parse_table(lines, index)
                    if parsed:
                        table_data, new_index = parsed
                        if table_data:
                            insert_table(document, table_data)
                            index = new_index
                            continue
                if stripped_line.startswith('>'):
                    quote_match = re.match(r'^(>+)\s*(.*)$', stripped_line)
                    if quote_match:
                        quote_level = len(quote_match.group(1))
                        quote_text = quote_match.group(2).strip()
                        paragraph = document.add_paragraph()
                        paragraph.paragraph_format.left_indent = Inches(0.25 * quote_level)
                        paragraph.paragraph_format.space_before = Pt(6)
                        paragraph.paragraph_format.space_after = Pt(6)
                        pPr = paragraph._p.get_or_add_pPr()
                        pBdr = OxmlElement('w:pBdr')
                        for i in range(quote_level):
                            left = OxmlElement('w:left')
                            left.set(qn('w:val'), 'single')
                            left.set(qn('w:sz'), str(6 + (i * 2)))
                            left.set(qn('w:space'), str(20 * i))
                            left.set(qn('w:color'), '666666')
                            pBdr.append(left)
                        pPr.append(pBdr)
                        for run in paragraph.runs:
                            run.italic = True
                            run.font.color.rgb = RGBColor(85, 85, 85)
                        parse_inline_formatting(paragraph, quote_text)
                    index += 1
                    continue
                if not stripped_line:
                    document.add_paragraph()
                    # Reset list counters when we encounter an empty line
                    list_counters = {}
                    index += 1
                    continue
                heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped_line)
                if heading_match:
                    level = len(heading_match.group(1))
                    heading_text = heading_match.group(2).strip()
                    paragraph = document.add_paragraph()
                    paragraph.style = f'Heading {level}'
                    set_heading_style(paragraph, level)
                    parse_inline_formatting(paragraph, heading_text)
                    anchor_name = generate_anchor_name(heading_text)
                    bookmark_id += 1
                    add_bookmark(paragraph, anchor_name, bookmark_id)
                    # Reset list counters when we encounter a heading
                    list_counters = {}
                    index += 1
                    continue
                list_match = re.match(r'^(\s*)([-*+]|(\d+\.))\s+(.*)', line)
                if list_match:
                    indent = get_list_level(line)
                    handle_list_item(document, line, indent, list_counters)
                    index += 1
                    continue
                paragraph = document.add_paragraph()
                parse_inline_formatting(paragraph, stripped_line)
                # Reset list counters when we encounter a regular paragraph
                list_counters = {}
                index += 1
        if in_code_block and code_block_content:
            code_text = '\n'.join(code_block_content)
            if code_block_language == 'chart':
                image_bytes = process_chart_code(code_text)
                insert_chart_image(document, image_bytes)
            elif code_block_language == 'mermaid':
                image_bytes = process_mermaid_diagram(code_text)
                insert_chart_image(document, image_bytes)
            else:
                add_code_block(document, code_text, code_block_language)

        document.save(file_path)
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        messagebox.showerror("Error", f"An error occurred: {e}")

def generate_anchor_name(heading_text):
    anchor = heading_text.lower()
    anchor = re.sub(r'[^a-z0-9\s\-]', '', anchor)
    anchor = re.sub(r'\s+', '-', anchor)
    return anchor

def add_page_break(document):
    """
    Add a page break to the Word document.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
