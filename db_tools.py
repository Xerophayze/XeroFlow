# db_tools.py

import subprocess
import sys
import logging
import os
import stat
import faiss
import numpy as np
from typing import List, Dict, Any, Optional
from contextlib import contextmanager
from datetime import datetime
import hashlib
import time
from filelock import FileLock
from langchain_community.document_loaders import PyPDFLoader, CSVLoader, TextLoader
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:  # pragma: no cover - compatibility shim
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document
import json
import shutil
import uuid  # For generating unique chunk IDs
import docx  # For .docx files
import tempfile

try:
    import torch
except ImportError:
    torch = None

# Try to import Windows-specific modules, but don't fail if they're not available
try:
    import win32com.client
    import pythoncom
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False
    logging.warning("win32com not available. Legacy .doc file support will be disabled.")

# Setup logging
LOG_DIR = "logs"
if not os.path.exists(LOG_DIR):
    os.makedirs(LOG_DIR)

logging.basicConfig(
    filename=os.path.join(LOG_DIR, "app.log"),
    filemode='a',
    format='%(asctime)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

DATABASES_DIR = "databases"
METADATA_FILE = "metadata.json"
DOCUMENTS_INDEX_FILE = "documents.json"
NOTES_FILE = "notes.json"
FAISS_INDEX_FILE = "faiss.index"
LOCK_FILE_NAME = "db.lock"
SEARCH_METRICS_FILE = "search_metrics.jsonl"
DEFAULT_CHUNK_SIZE = 800
DEFAULT_CHUNK_OVERLAP = 150
DEFAULT_MAX_CONTEXT_CHUNKS = 3
DEFAULT_FILTERABLE_FIELDS = {"tags", "doc_id", "source"}
MMR_LAMBDA = 0.6

class DatabaseManager:
    def __init__(self):
        # Switch to a more robust model
        self.device = self._detect_device()
        model_kwargs = {"device": self.device} if self.device else {}
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L12-v2",
            model_kwargs=model_kwargs
        )
        self.embedding_dimension = self._get_embedding_dimension()
        if not os.path.exists(DATABASES_DIR):
            os.makedirs(DATABASES_DIR)
        logging.info("Initialized DatabaseManager (embedding device=%s).", self.device or "unknown")

    def _db_path(self, db_name: str) -> str:
        return os.path.join(DATABASES_DIR, db_name)

    def _ensure_db_dir(self, db_name: str) -> str:
        path = self._db_path(db_name)
        os.makedirs(path, exist_ok=True)
        return path

    def _metadata_path(self, db_name: str) -> str:
        return os.path.join(self._db_path(db_name), METADATA_FILE)

    def _documents_index_path(self, db_name: str) -> str:
        return os.path.join(self._db_path(db_name), DOCUMENTS_INDEX_FILE)

    def _notes_path(self, db_name: str) -> str:
        return os.path.join(self._db_path(db_name), NOTES_FILE)

    def _faiss_path(self, db_name: str) -> str:
        return os.path.join(self._db_path(db_name), FAISS_INDEX_FILE)

    def _lock_path(self, db_name: str) -> str:
        return os.path.join(self._db_path(db_name), LOCK_FILE_NAME)

    @contextmanager
    def _with_db_lock(self, db_name: str):
        self._ensure_db_dir(db_name)
        lock = FileLock(self._lock_path(db_name))
        lock.acquire()
        try:
            yield
        finally:
            lock.release()

    def _read_json_file(self, path: str, default):
        if not os.path.exists(path):
            return default
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as exc:
            logging.error(f"Failed to read JSON at {path}: {exc}")
            return default

    def _write_json_file(self, path: str, payload):
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(payload, f, indent=2)
        except Exception as exc:
            logging.error(f"Failed to write JSON at {path}: {exc}")

    def _calculate_file_hash(self, file_path: str) -> str:
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _get_text_splitter(self, chunk_size: int = DEFAULT_CHUNK_SIZE, chunk_overlap: int = DEFAULT_CHUNK_OVERLAP):
        return RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )

    def _embed_text(self, text: str) -> Optional[np.ndarray]:
        if not text:
            return None
        embedding = self.embeddings.embed_documents([text])[0]
        embedding = np.array(embedding).astype('float32')
        norm = np.linalg.norm(embedding)
        if norm == 0:
            logging.warning("Encountered zero norm embedding.")
            return None
        embedding /= norm
        return embedding

    def _get_embedding_dimension(self) -> int:
        test_embedding = self.embeddings.embed_query("dimension_probe")
        embedding = np.array(test_embedding).astype('float32')
        if embedding.ndim != 1:
            embedding = embedding.flatten()
        return embedding.shape[0]

    def _load_documents_index(self, db_name: str) -> List[Dict[str, Any]]:
        return self._read_json_file(self._documents_index_path(db_name), [])

    def _save_documents_index(self, db_name: str, documents: List[Dict[str, Any]]):
        self._write_json_file(self._documents_index_path(db_name), documents)

    def _load_chunk_metadata(self, db_name: str) -> List[Dict[str, Any]]:
        return self._read_json_file(self._metadata_path(db_name), [])

    def _save_chunk_metadata(self, db_name: str, metadata: List[Dict[str, Any]]):
        self._write_json_file(self._metadata_path(db_name), metadata)

    def _ensure_faiss_index(self, db_name: str) -> faiss.Index:
        index = self.load_faiss_index(db_name)
        if index is None:
            index = faiss.IndexFlatIP(self.embedding_dimension)
            self.save_faiss_index(db_name, index)
        return index

    def _rebuild_faiss_index(self, db_name: str, metadata: List[Dict[str, Any]]):
        index = faiss.IndexFlatIP(self.embedding_dimension)
        vectors = []
        for entry in metadata:
            vector = entry.get("embedding")
            if vector is None:
                vector = self._embed_text(entry.get("content", ""))
                if vector is None:
                    continue
                entry["embedding"] = vector.tolist()
            vector_np = np.array(vector, dtype='float32')
            norm = np.linalg.norm(vector_np)
            if norm == 0:
                continue
            vector_np /= norm
            entry["embedding"] = vector_np.tolist()
            vectors.append(vector_np)
        if vectors:
            stacked = np.vstack(vectors).astype('float32')
            index.add(stacked)
        self.save_faiss_index(db_name, index)

    def _matches_filters(self, chunk: Dict[str, Any], doc_lookup: Dict[str, Dict[str, Any]], filters: Dict[str, Any]) -> bool:
        if not filters:
            return True
        doc_meta = doc_lookup.get(chunk.get("doc_id"), {})
        for key, expected in filters.items():
            if key == "tags":
                doc_tags = set(doc_meta.get("tags", []))
                expected_set = set(expected) if isinstance(expected, (list, tuple, set)) else {expected}
                if not expected_set.issubset(doc_tags):
                    return False
            elif key == "doc_id":
                if chunk.get("doc_id") != expected:
                    return False
            elif key == "source":
                if chunk.get("source") != expected:
                    return False
            else:
                chunk_value = chunk.get(key) or doc_meta.get(key)
                if isinstance(expected, (list, tuple, set)):
                    if chunk_value not in expected:
                        return False
                else:
                    if chunk_value != expected:
                        return False
        return True

    def _apply_mmr(self, query_embedding: np.ndarray, candidates: List[Dict[str, Any]], top_k: int) -> List[Dict[str, Any]]:
        if not candidates:
            return []
        candidate_embeddings = [
            np.array(c["embedding"], dtype='float32')
            for c in candidates
            if c.get("embedding") is not None
        ]
        if not candidate_embeddings:
            return candidates[:top_k]
        candidate_embeddings = [vec / (np.linalg.norm(vec) or 1.0) for vec in candidate_embeddings]
        query = query_embedding / (np.linalg.norm(query_embedding) or 1.0)
        selected = []
        selected_indices = []
        available_indices = list(range(len(candidate_embeddings)))
        while available_indices and len(selected) < top_k:
            mmr_scores = []
            for idx in available_indices:
                candidate_vec = candidate_embeddings[idx]
                similarity_to_query = float(np.dot(candidate_vec, query))
                if not selected_indices:
                    diversity = 0.0
                else:
                    diversity = max(float(np.dot(candidate_vec, candidate_embeddings[s_idx])) for s_idx in selected_indices)
                score = MMR_LAMBDA * similarity_to_query - (1 - MMR_LAMBDA) * diversity
                mmr_scores.append((score, idx))
            _, best_idx = max(mmr_scores, key=lambda item: item[0])
            selected.append(candidates[best_idx])
            selected_indices.append(best_idx)
            available_indices.remove(best_idx)
        return selected

    def _assemble_context(self, chunk: Dict[str, Any], doc_chunks: List[Dict[str, Any]], window: int = DEFAULT_MAX_CONTEXT_CHUNKS) -> str:
        if not doc_chunks:
            return chunk.get("content", "")
        sorted_chunks = sorted(doc_chunks, key=lambda c: c.get("chunk_number", 0))
        target_index = next((i for i, entry in enumerate(sorted_chunks) if entry.get("chunk_id") == chunk.get("chunk_id")), None)
        if target_index is None:
            return chunk.get("content", "")
        start = max(0, target_index - window)
        end = min(len(sorted_chunks), target_index + window + 1)
        combined = []
        for entry in sorted_chunks[start:end]:
            heading = entry.get("section") or entry.get("metadata", {}).get("heading")
            prefix = f"[{heading}]\n" if heading else ""
            combined.append(prefix + entry.get("content", ""))
        return "\n\n".join(part for part in combined if part)

    def _log_search_metric(self, db_name: str, payload: Dict[str, Any]):
        record = payload.copy()
        record["timestamp"] = datetime.utcnow().isoformat()
        metrics_path = os.path.join(self._db_path(db_name), SEARCH_METRICS_FILE)
        try:
            with open(metrics_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(record) + "\n")
        except Exception as exc:
            logging.error(f"Failed to log search metrics: {exc}")

    @staticmethod
    def _handle_remove_readonly(func, path, exc_info):
        try:
            os.chmod(path, stat.S_IWRITE)
            func(path)
        except Exception as exc:  # pragma: no cover - best-effort cleanup
            logging.error("Failed to remove path '%s': %s", path, exc)

    def _detect_device(self) -> str:
        """Determine the preferred device for embedding computations."""
        if torch is None:
            logging.warning("PyTorch not available. Defaulting embeddings to CPU.")
            return "cpu"

        if torch.cuda.is_available():
            try:
                device_index = torch.cuda.current_device()
                device_name = torch.cuda.get_device_name(device_index)
                capability = torch.cuda.get_device_capability(device_index)
                logging.info(
                    "CUDA available - using GPU %s: %s (compute capability %s.%s)",
                    device_index,
                    device_name,
                    capability[0],
                    capability[1]
                )
            except Exception as cuda_error:
                logging.warning("Failed to query CUDA device details: %s", cuda_error)
            return "cuda"

        logging.info("CUDA not available. Falling back to CPU for embeddings.")
        return "cpu"
    
    def preprocess_content(self, content: str) -> str:
        """
        Preprocess the content by removing irrelevant sections.
        
        Args:
            content (str): The raw text content.
        
        Returns:
            str: The cleaned text content.
        """
        # Example: Remove lines containing certain keywords
        keywords = ['page', 'copyright', 'confidential', 'scan the qr code', 'seats', 'policies', 'help']
        lines = content.split('\n')
        filtered_lines = [line for line in lines if not any(keyword in line.lower() for keyword in keywords)]
        return '\n'.join(filtered_lines)
    
    def _get_file_type(self, uploaded_file: Any) -> str:
        """
        Determine the file type based on the file extension.

        Args:
            uploaded_file (file object): File object opened in binary mode.

        Returns:
            str: The file type ('pdf', 'csv', 'txt', 'docx', 'doc', etc.).
        """
        _, ext = os.path.splitext(uploaded_file.name)
        file_type = ext.lower().lstrip('.')
        logging.debug(f"Determined file type '{file_type}' for file '{uploaded_file.name}'.")
        return file_type

    def _save_uploaded_file(self, db_name: str, uploaded_file: Any) -> str:
        """
        Save the uploaded file to the specified database directory.

        Args:
            db_name (str): Name of the database.
            uploaded_file (file object): File object opened in binary mode.

        Returns:
            str: Path to the saved file.
        """
        db_path = self._ensure_db_dir(db_name)
        file_name = os.path.basename(uploaded_file.name)
        file_path = os.path.abspath(os.path.join(db_path, file_name))
        try:
            try:
                uploaded_file.seek(0)
            except Exception:
                pass
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            logging.info(f"Saved uploaded file '{uploaded_file.name}' to '{file_path}'.")
            return file_path
        except Exception as e:
            logging.error(f"Failed to save uploaded file '{uploaded_file.name}': {e}")
            raise e

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from a .docx file.
        
        Args:
            file_path (str): Path to the .docx file.
            
        Returns:
            str: Extracted text content.
        """
        try:
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # Skip empty cells
                            row_text.append(cell.text)
                    if row_text:  # Skip empty rows
                        text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            logging.error(f"Error extracting text from .docx file '{file_path}': {e}")
            raise

    def _extract_text_from_doc(self, file_path: str) -> str:
        """
        Extract text from a legacy .doc file using Word COM object.
        
        Args:
            file_path (str): Path to the .doc file.
            
        Returns:
            str: Extracted text content.
        """
        if not WORD_SUPPORT:
            raise ImportError("win32com is not available. Cannot process .doc files.")
            
        try:
            pythoncom.CoInitialize()  # Initialize COM
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Create a temporary file for the .docx version
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = temp_file.name
            
            try:
                # Open and save as .docx
                doc = word.Documents.Open(file_path)
                doc.SaveAs2(temp_path, FileFormat=16)  # 16 = .docx format
                doc.Close()
                
                # Extract text from the .docx version
                text = self._extract_text_from_docx(temp_path)
                
                return text
            finally:
                word.Quit()
                # Clean up the temporary file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                pythoncom.CoUninitialize()
        except Exception as e:
            logging.error(f"Error extracting text from .doc file '{file_path}': {e}")
            raise

    def _load_documents_from_path(self, file_type: str, file_path: str, source_name: str) -> List[Document]:
        try:
            if file_type == 'pdf':
                loader = PyPDFLoader(file_path)
                docs = loader.load()
            elif file_type == 'csv':
                loader = CSVLoader(file_path)
                docs = loader.load()
            elif file_type == 'txt':
                loader = TextLoader(file_path)
                docs = loader.load()
            elif file_type == 'docx':
                text = self._extract_text_from_docx(file_path)
                docs = [Document(page_content=text, metadata={"source": source_name})]
            elif file_type == 'doc':
                if not WORD_SUPPORT:
                    logging.warning("Skipping .doc file '%s' - win32com not available", source_name)
                    return []
                text = self._extract_text_from_doc(file_path)
                docs = [Document(page_content=text, metadata={"source": source_name})]
            else:
                logging.warning("Unsupported file type '%s' for '%s'", file_type, source_name)
                return []

            normalized_docs = []
            for doc in docs:
                metadata = dict(doc.metadata or {})
                metadata.setdefault("source", source_name)
                normalized_docs.append(Document(page_content=doc.page_content, metadata=metadata))
            return normalized_docs
        except Exception as exc:
            logging.error("Failed to load '%s': %s", source_name, exc)
            return []

    def add_documents(self, db_name: str, uploaded_files: List[Any]) -> Dict[str, Any]:
        """Add or update documents in the specified database with rich metadata tracking."""
        summary = {"added": [], "updated": [], "skipped": []}
        try:
            with self._with_db_lock(db_name):
                index = self._ensure_faiss_index(db_name)
                documents_index = self._load_documents_index(db_name)
                chunk_metadata = self._load_chunk_metadata(db_name)
                documents_by_source = {doc.get("source"): doc for doc in documents_index}

                needs_rebuild = False
                vectors_to_add: List[np.ndarray] = []

                for uploaded_file in uploaded_files:
                    file_name = os.path.basename(uploaded_file.name)
                    file_path = self._save_uploaded_file(db_name, uploaded_file)
                    file_type = self._get_file_type(uploaded_file)
                    file_hash = self._calculate_file_hash(file_path)
                    timestamp = datetime.utcnow().isoformat()

                    existing_doc = documents_by_source.get(file_name)
                    if existing_doc and existing_doc.get("file_hash") == file_hash:
                        logging.info("File '%s' unchanged (hash match). Skipping.", file_name)
                        summary["skipped"].append(file_name)
                        continue

                    documents = self._load_documents_from_path(file_type, file_path, file_name)
                    if not documents:
                        logging.warning("No documents parsed from '%s'. Skipping.", file_name)
                        summary["skipped"].append(file_name)
                        continue

                    splitter = self._get_text_splitter()
                    chunks = splitter.split_documents(documents)
                    if not chunks:
                        summary["skipped"].append(file_name)
                        continue

                    doc_id = existing_doc.get("doc_id") if existing_doc else str(uuid.uuid4())
                    if existing_doc:
                        chunk_metadata = [chunk for chunk in chunk_metadata if chunk.get("doc_id") != doc_id]
                        needs_rebuild = True

                    new_chunk_entries = []
                    chunk_number = 0
                    for chunk in chunks:
                        cleaned_content = self.preprocess_content(chunk.page_content)
                        if not cleaned_content.strip():
                            continue
                        embedding = self._embed_text(cleaned_content)
                        if embedding is None:
                            continue
                        chunk_entry = {
                            "chunk_id": str(uuid.uuid4()),
                            "doc_id": doc_id,
                            "source": file_name,
                            "chunk_number": chunk_number,
                            "page": chunk.metadata.get("page"),
                            "section": chunk.metadata.get("section"),
                            "content": cleaned_content,
                            "embedding": embedding.tolist(),
                            "created_at": timestamp
                        }
                        new_chunk_entries.append(chunk_entry)
                        if not needs_rebuild and not existing_doc:
                            vectors_to_add.append(embedding)
                        chunk_number += 1

                    if not new_chunk_entries:
                        summary["skipped"].append(file_name)
                        continue

                    chunk_metadata.extend(new_chunk_entries)

                    doc_record = existing_doc or {
                        "doc_id": doc_id,
                        "source": file_name,
                        "tags": [],
                        "metadata": {}
                    }
                    doc_record.update({
                        "path": file_path,
                        "file_type": file_type,
                        "file_hash": file_hash,
                        "size_bytes": os.path.getsize(file_path),
                        "added_at": doc_record.get("added_at", timestamp),
                        "updated_at": timestamp,
                        "chunk_count": len(new_chunk_entries)
                    })
                    if not existing_doc:
                        documents_index.append(doc_record)
                        documents_by_source[file_name] = doc_record
                        summary["added"].append(file_name)
                    else:
                        summary["updated"].append(file_name)

                self._save_chunk_metadata(db_name, chunk_metadata)
                self._save_documents_index(db_name, documents_index)

                if needs_rebuild:
                    self._rebuild_faiss_index(db_name, chunk_metadata)
                elif vectors_to_add:
                    stacked = np.vstack(vectors_to_add).astype('float32')
                    index.add(stacked)
                    self.save_faiss_index(db_name, index)

            logging.info("Document ingest summary for '%s': %s", db_name, summary)
            return {"success": True, "summary": summary}
        except Exception as exc:
            error_msg = f"Error adding documents: {exc}"
            logging.exception(error_msg)
            return {"success": False, "error": error_msg}

    def create_database(self, db_name: str) -> bool:
        """Create a new database directory with empty indexes and metadata files."""
        db_path = self._db_path(db_name)
        if os.path.exists(db_path):
            logging.warning("Database '%s' already exists.", db_name)
            return False
        try:
            os.makedirs(db_path, exist_ok=True)
            index = faiss.IndexFlatIP(self.embedding_dimension)
            self.save_faiss_index(db_name, index)
            self._save_chunk_metadata(db_name, [])
            self._save_documents_index(db_name, [])
            self.save_notes(db_name, {})
            metrics_path = os.path.join(db_path, SEARCH_METRICS_FILE)
            open(metrics_path, "a", encoding="utf-8").close()
            logging.info("Initialized database '%s' with dimension %s.", db_name, self.embedding_dimension)
            return True
        except Exception as exc:
            logging.error("Failed to create database '%s': %s", db_name, exc)
            shutil.rmtree(db_path, ignore_errors=True)
            return False

    def delete_document(self, db_name: str, doc_identifier: str) -> Dict[str, Any]:
        """Remove a document (by name or path) and rebuild indexes."""
        try:
            with self._with_db_lock(db_name):
                documents = self._load_documents_index(db_name)
                chunk_metadata = self._load_chunk_metadata(db_name)
                target_name = os.path.basename(doc_identifier)
                doc_record = next((doc for doc in documents if doc.get("source") == target_name), None)
                if not doc_record:
                    return {"success": False, "error": f"Document '{target_name}' not found."}

                doc_id = doc_record.get("doc_id")
                documents = [doc for doc in documents if doc.get("doc_id") != doc_id]
                chunk_metadata = [chunk for chunk in chunk_metadata if chunk.get("doc_id") != doc_id]

                self._save_documents_index(db_name, documents)
                self._save_chunk_metadata(db_name, chunk_metadata)
                self._rebuild_faiss_index(db_name, chunk_metadata)

                file_path = doc_record.get("path") or os.path.join(self._db_path(db_name), target_name)
                if file_path and os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as exc:
                        logging.warning("Failed to remove file '%s': %s", file_path, exc)

            logging.info("Deleted document '%s' from '%s'.", target_name, db_name)
            return {"success": True}
        except Exception as exc:
            error_msg = f"Failed to delete document: {exc}"
            logging.error(error_msg)
            return {"success": False, "error": error_msg}

    def delete_database(self, db_name: str) -> Dict[str, Any]:
        """
        Delete an entire database, including all its documents, metadata, and indexes.

        Args:
            db_name (str): Name of the database to delete.

        Returns:
            Dict[str, Any]: Result of the operation.
        """
        try:
            db_path = self._db_path(db_name)
            if not os.path.exists(db_path):
                error_msg = f"Database '{db_name}' does not exist."
                logging.error(error_msg)
                return {"success": False, "error": error_msg}

            lock_path = self._lock_path(db_name)
            lock = FileLock(lock_path)
            # Acquire lock to ensure no concurrent writers hold open handles during deletion
            lock.acquire()
            try:
                # Release happens in finally to ensure deterministic unlock
                pass
            finally:
                lock.release()

            # Remove leftover lock file to avoid access denied on Windows
            if os.path.exists(lock_path):
                try:
                    os.remove(lock_path)
                except Exception as exc:
                    logging.warning("Failed to remove lock file '%s': %s", lock_path, exc)

            shutil.rmtree(db_path)
            logging.info(f"Deleted database '{db_name}' successfully.")
            return {"success": True}
        except Exception as e:
            error_msg = f"Failed to delete database '{db_name}': {e}"
            logging.error(error_msg)
            return {"success": False, "error": error_msg}

    def list_databases(self) -> List[str]:
        """
        List all available databases.

        Returns:
            List[str]: List of database names.
        """
        try:
            dbs = [name for name in os.listdir(DATABASES_DIR)
                   if os.path.isdir(os.path.join(DATABASES_DIR, name))]
            logging.info(f"Listed databases: {dbs}")
            return dbs
        except Exception as e:
            logging.error(f"Failed to list databases: {e}")
            return []

    def list_documents(self, db_name: str) -> List[str]:
        """Return the list of stored document names."""
        try:
            docs = self.list_document_records(db_name)
            names = [doc.get("source") for doc in docs]
            logging.info("Listed %d documents for '%s'.", len(names), db_name)
            return names
        except Exception as exc:
            logging.error("Error listing documents for '%s': %s", db_name, exc)
            return []

    def list_document_records(self, db_name: str) -> List[Dict[str, Any]]:
        """Return detailed document metadata records for a database."""
        try:
            records = self._load_documents_index(db_name)
            logging.info("Loaded %d document records for '%s'.", len(records), db_name)
            return records
        except Exception as exc:
            logging.error("Error loading document records for '%s': %s", db_name, exc)
            return []

    def load_faiss_index(self, db_name: str) -> faiss.Index:
        """
        Load the FAISS index for the specified database.

        Args:
            db_name (str): Name of the database.

        Returns:
            faiss.Index: Loaded FAISS index.
        """
        index_path = self._faiss_path(db_name)
        if not os.path.exists(index_path):
            logging.warning(f"FAISS index not found for database '{db_name}'.")
            return None
        try:
            index = faiss.read_index(index_path)
            logging.info(f"Loaded FAISS index for database '{db_name}'.")
            return index
        except Exception as e:
            logging.error(f"Failed to load FAISS index for database '{db_name}': {e}")
            return None

    def save_faiss_index(self, db_name: str, index: faiss.Index):
        """
        Save the FAISS index for the specified database.

        Args:
            db_name (str): Name of the database.
            index (faiss.Index): FAISS index to save.
        """
        index_path = os.path.join(DATABASES_DIR, db_name, "faiss.index")
        try:
            faiss.write_index(index, index_path)
            logging.info(f"Saved FAISS index for database '{db_name}'.")
        except Exception as e:
            logging.error(f"Failed to save FAISS index for database '{db_name}': {e}")

    def load_metadata(self, db_name: str) -> List[Dict[str, Any]]:
        """
        Load metadata for the specified database.

        Args:
            db_name (str): Name of the database.

        Returns:
            List[Dict[str, Any]]: List of metadata dictionaries.
        """
        return self._load_chunk_metadata(db_name)

    def save_metadata(self, db_name: str, metadata: List[Dict[str, Any]]):
        """
        Save metadata for the specified database.

        Args:
            db_name (str): Name of the database.
            metadata (List[Dict[str, Any]]): List of metadata dictionaries.
        """
        self._save_chunk_metadata(db_name, metadata)

    def load_notes(self, db_name: str) -> Dict[str, str]:
        """
        Load notes for the specified database.

        Args:
            db_name (str): Name of the database.

        Returns:
            Dict[str, str]: Dictionary of notes.
        """
        return self._read_json_file(self._notes_path(db_name), {})

    def save_notes(self, db_name: str, notes: Dict[str, str]):
        """
        Save notes for the specified database.

        Args:
            db_name (str): Name of the database.
            notes (Dict[str, str]): Dictionary of notes.
        """
        self._write_json_file(self._notes_path(db_name), notes)

    def search(
        self,
        db_name: str,
        query: str,
        top_k: int = 10,
        filters: Optional[Dict[str, Any]] = None,
        context_window: int = DEFAULT_MAX_CONTEXT_CHUNKS,
        rerank: bool = True,
        collect_metrics: bool = True
    ) -> List[Dict[str, Any]]:
        """Search the database with optional metadata filters and reranking."""
        start_time = time.time()
        filters = filters or {}
        try:
            index = self.load_faiss_index(db_name)
            if index is None or index.ntotal == 0:
                logging.warning("FAISS index unavailable or empty for '%s'.", db_name)
                return []

            chunk_metadata = self._load_chunk_metadata(db_name)
            if not chunk_metadata:
                logging.info("No chunk metadata found for '%s'.", db_name)
                return []

            documents = self._load_documents_index(db_name)
            doc_lookup = {doc.get("doc_id"): doc for doc in documents}

            query_embedding = self.embeddings.embed_query(query)
            query_embedding = np.array(query_embedding, dtype='float32')
            norm = np.linalg.norm(query_embedding)
            if norm == 0:
                logging.warning("Zero-norm query embedding for '%s'.", query)
                return []
            query_embedding /= norm

            fetch_k = min(max(top_k * 4, top_k), len(chunk_metadata))
            D, I = index.search(np.array([query_embedding]).astype('float32'), fetch_k)

            candidates = []
            seen_chunk_ids = set()
            for score, idx in zip(D[0], I[0]):
                if idx < 0 or idx >= len(chunk_metadata):
                    continue
                chunk = chunk_metadata[idx]
                chunk_id = chunk.get("chunk_id")
                if not chunk_id or chunk_id in seen_chunk_ids:
                    continue
                if not self._matches_filters(chunk, doc_lookup, filters):
                    continue

                embedding = np.array(chunk.get("embedding"), dtype='float32') if chunk.get("embedding") is not None else None
                if embedding is None or embedding.size == 0:
                    embedding = self._embed_text(chunk.get("content", ""))
                    if embedding is None:
                        continue
                    chunk["embedding"] = embedding.tolist()

                chunk_copy = chunk.copy()
                chunk_copy["similarity"] = float(score)
                chunk_copy["embedding"] = embedding
                candidates.append(chunk_copy)
                seen_chunk_ids.add(chunk_id)
                if len(candidates) >= fetch_k:
                    break

            if not candidates:
                logging.info("No candidates found for query '%s' in '%s'.", query, db_name)
                return []

            if rerank and len(candidates) > top_k:
                ranked_chunks = self._apply_mmr(query_embedding, candidates, top_k)
            else:
                ranked_chunks = candidates[:top_k]

            doc_chunks_map: Dict[str, List[Dict[str, Any]]] = {}
            for chunk in chunk_metadata:
                doc_chunks_map.setdefault(chunk.get("doc_id"), []).append(chunk)

            results = []
            for chunk in ranked_chunks:
                doc_id = chunk.get("doc_id")
                doc_chunks = doc_chunks_map.get(doc_id, [])
                context_text = self._assemble_context(chunk, doc_chunks, window=context_window)
                results.append({
                    "doc_id": doc_id,
                    "source": chunk.get("source"),
                    "similarity": chunk.get("similarity", 0.0),
                    "content": context_text,
                    "chunk_id": chunk.get("chunk_id"),
                    "metadata": {
                        "page": chunk.get("page"),
                        "chunk_number": chunk.get("chunk_number"),
                        "section": chunk.get("section"),
                        "filters": filters,
                    },
                    "document": doc_lookup.get(doc_id, {}),
                })

            if collect_metrics:
                latency_ms = (time.time() - start_time) * 1000
                self._log_search_metric(db_name, {
                    "query": query,
                    "top_k": top_k,
                    "filters": filters,
                    "results": len(results),
                    "latency_ms": latency_ms
                })

            logging.info("Search completed for '%s' (%d results).", db_name, len(results))
            return results
        except Exception as exc:
            logging.error("Error during search: %s", exc)
            return []
    
    def add_note(self, db_name: str, chunk_id: str, note: str):
        """
        Add or update a note for a specific chunk.

        Args:
            db_name (str): Name of the database.
            chunk_id (str): ID of the chunk.
            note (str): Note content.
        """
        try:
            notes = self.load_notes(db_name)
            notes[chunk_id] = note
            self.save_notes(db_name, notes)
            logging.info(f"Added/Updated note for chunk ID {chunk_id} in database '{db_name}'.")
        except Exception as e:
            logging.error(f"Error adding note: {e}")

    def get_note(self, db_name: str, chunk_id: str) -> str:
        """
        Retrieve a note for a specific chunk.

        Args:
            db_name (str): Name of the database.
            chunk_id (str): ID of the chunk.

        Returns:
            str: Note content.
        """
        try:
            notes = self.load_notes(db_name)
            note = notes.get(chunk_id, "")
            logging.info(f"Retrieved note for chunk ID {chunk_id} in database '{db_name}'.")
            return note
        except Exception as e:
            logging.error(f"Error getting note: {e}")
            return ""
