# nodes/Assistant_node.py
import sys
import subprocess
try:
    from importlib.metadata import distribution, PackageNotFoundError
except ImportError:  # Python <3.8 fallback if needed
    from importlib_metadata import distribution, PackageNotFoundError
import tkinter as tk
from tkinter import ttk, scrolledtext
from threading import Thread, Event
import os
import json
import time
import shutil
import requests
import re
from pathlib import Path
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from .base_node import BaseNode
from src.workflows.node_registry import register_node
from src.api.handler import process_api_request
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .youtube_transcript_node import YoutubeTranscriptNode
from .web_scrape_node import WebScrapingNode
from src.export.word import convert_markdown_to_docx
from src.export.process_output import process_api_output
from pydub import AudioSegment
import traceback
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import os
import json
import time
from datetime import datetime
import docx
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse
import pandas as pd
import PyPDF2
import io
import difflib

@register_node('AssistantNode')
class AssistantNode(BaseNode):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.stop_event = Event()
        self.processed_files = []
        self.monitor_window = None
        self.output_text = None
        self.current_job_folder = None  # Store the current job's output folder
        self.custom_project_name = None  # Store the custom project name if provided

    def define_inputs(self):
        return []  # No inputs needed as this node monitors a folder

    def define_outputs(self):
        return ['output']  # Output will be the API response

    def define_properties(self):
        props = self.get_default_properties()
        
        # Get the options lists upfront
        api_endpoints = self.get_api_endpoints()
        search_endpoints = self.get_search_api_endpoints()
        
        props.update({
            'node_name': {
                'type': 'text',
                'label': 'Custom Node Name',
                'default': 'AssistantNode'
            },
            'description': {
                'type': 'text',
                'label': 'Description',
                'default': 'Monitors a folder for text and audio files, processes them through respective APIs.'
            },
            'inbox_folder': {
                'type': 'text',
                'label': 'Inbox Folder Path',
                'default': ''
            },
            'outbox_folder': {
                'type': 'text',
                'label': 'Outbox Folder Path',
                'default': ''
            },
            'text_api_endpoint': {
                'type': 'dropdown',
                'label': 'Text Processing API',
                'options': api_endpoints,  # Use the actual list instead of function reference
                'default': api_endpoints[0] if api_endpoints else ''
            },
            'whisper_api_endpoint': {
                'type': 'dropdown',
                'label': 'Audio Transcription API',
                'options': api_endpoints,  # Use the actual list instead of function reference
                'default': api_endpoints[0] if api_endpoints else ''
            },
            'search_api_endpoint': {
                'type': 'dropdown',
                'label': 'Search API',
                'options': search_endpoints,  # Use the actual list instead of function reference
                'default': search_endpoints[0] if search_endpoints else ''
            },
            'search_url': {
                'type': 'text',
                'label': 'Search URL (Deprecated)',
                'default': ''
            },
            'num_search_results': {
                'type': 'number',
                'label': 'Number of Search Results',
                'default': 3,
                'min': 1,
                'max': 10
            },
            'Prompt': {
                'type': 'textarea',
                'label': 'Prompt Template',
                'default': ''
            },
            'PreProcess': {
                'type': 'textarea',
                'label': 'Pre-Process Text',
                'default': ''
            }
        })
        return props

    def get_api_endpoints(self):
        """Get available API endpoints with dynamic refresh"""
        interfaces = self.config.get('interfaces', {})
        if interfaces is None:
            interfaces = {}
        api_list = list(interfaces.keys())
        return api_list

    def get_search_api_endpoints(self):
        """Get list of search API endpoints."""
        return [api for api, config in self.config.get('interfaces', {}).items() 
                if config.get('type', '').lower() == 'searchengine']

    def create_monitor_window(self):
        """Create a window to display monitoring status and project submission interface."""
        if not hasattr(self, 'monitor_window') or not self.monitor_window:
            self.monitor_window = tk.Toplevel()
            
            # Get workflow name from inputs if available
            window_title = "Assistant Interface"
            if hasattr(self, 'inputs') and isinstance(self.inputs, dict):
                # Try to get workflow name from workflow_id
                if 'workflow_id' in self.inputs:
                    from main import workflow_manager
                    workflow_id = self.inputs.get('workflow_id')
                    workflow = workflow_manager.get_workflow(workflow_id)
                    if workflow and workflow.workflow_name:
                        window_title = f"Assistant Interface - {workflow.workflow_name}"
                # Or directly from workflow_name if present
                elif 'workflow_name' in self.inputs:
                    workflow_name = self.inputs.get('workflow_name')
                    if workflow_name:
                        window_title = f"Assistant Interface - {workflow_name}"
            
            self.monitor_window.title(window_title)
            self.monitor_window.protocol("WM_DELETE_WINDOW", self.stop_monitoring)
            
            # Set window size and make it resizable
            self.monitor_window.geometry("900x700")
            self.monitor_window.minsize(700, 600)
            
            # Configure grid weights
            self.monitor_window.grid_rowconfigure(0, weight=1)
            self.monitor_window.grid_columnconfigure(0, weight=1)
            
            # Create a notebook for tabs
            notebook = ttk.Notebook(self.monitor_window)
            notebook.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
            
            # Create tabs
            submit_tab = ttk.Frame(notebook)
            output_tab = ttk.Frame(notebook)
            log_tab = ttk.Frame(notebook)
            
            notebook.add(submit_tab, text="Submit Project")
            notebook.add(output_tab, text="Output Files")
            notebook.add(log_tab, text="Processing Log")
            
            # Configure grid weights for tabs
            submit_tab.grid_rowconfigure(0, weight=0)  # Project name label
            submit_tab.grid_rowconfigure(1, weight=0)  # Project name entry
            submit_tab.grid_rowconfigure(2, weight=0)  # Instructions label
            submit_tab.grid_rowconfigure(3, weight=2)  # Instructions text
            submit_tab.grid_rowconfigure(4, weight=0)  # Files label
            submit_tab.grid_rowconfigure(5, weight=1)  # Files list
            submit_tab.grid_rowconfigure(6, weight=0)  # Buttons
            submit_tab.grid_columnconfigure(0, weight=1)
            
            output_tab.grid_rowconfigure(0, weight=0)  # Path frame
            output_tab.grid_rowconfigure(1, weight=1)  # File browser frame
            output_tab.grid_columnconfigure(0, weight=1)
            
            log_tab.grid_rowconfigure(0, weight=1)
            log_tab.grid_columnconfigure(0, weight=1)
            
            # === SUBMIT PROJECT TAB ===
            
            # Project name section
            project_name_label = ttk.Label(submit_tab, text="Project Name (optional):")
            project_name_label.grid(row=0, column=0, sticky="w", padx=5, pady=(5, 0))
            
            self.project_name_var = tk.StringVar()
            self.project_name_entry = ttk.Entry(submit_tab, textvariable=self.project_name_var, width=80)
            self.project_name_entry.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
            
            # Instructions section
            instructions_label = ttk.Label(submit_tab, text="Enter your instructions:")
            instructions_label.grid(row=2, column=0, sticky="w", padx=5, pady=(5, 0))
            
            self.instructions_text = scrolledtext.ScrolledText(submit_tab, height=10, width=80, wrap=tk.WORD)
            self.instructions_text.grid(row=3, column=0, sticky="nsew", padx=5, pady=5)
            
            # Files section
            files_frame = ttk.Frame(submit_tab)
            files_frame.grid(row=4, column=0, sticky="ew", padx=5, pady=(10, 0))
            
            files_label = ttk.Label(files_frame, text="Add files to your project:")
            files_label.pack(side=tk.LEFT, anchor="w")
            
            # Store the list of files to be included in the project
            self.project_files = []
            
            # Files list with scrollbar
            files_frame = ttk.Frame(submit_tab)
            files_frame.grid(row=5, column=0, sticky="nsew", padx=5, pady=5)
            
            self.files_listbox = tk.Listbox(files_frame, height=6, width=80)
            self.files_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            
            files_scrollbar = ttk.Scrollbar(files_frame, command=self.files_listbox.yview)
            files_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            self.files_listbox.config(yscrollcommand=files_scrollbar.set)
            
            # Buttons for file operations and submission
            buttons_frame = ttk.Frame(submit_tab)
            buttons_frame.grid(row=6, column=0, sticky="ew", padx=5, pady=10)
            
            add_file_button = ttk.Button(buttons_frame, text="Add Files", command=self.add_files)
            add_file_button.pack(side=tk.LEFT, padx=5)
            
            remove_file_button = ttk.Button(buttons_frame, text="Remove Selected", command=self.remove_selected_file)
            remove_file_button.pack(side=tk.LEFT, padx=5)
            
            clear_files_button = ttk.Button(buttons_frame, text="Clear All Files", command=self.clear_files)
            clear_files_button.pack(side=tk.LEFT, padx=5)
            
            # Add some space between file buttons and submit buttons
            ttk.Separator(buttons_frame, orient='vertical').pack(side=tk.LEFT, fill='y', padx=10, pady=5)
            
            submit_button = ttk.Button(buttons_frame, text="Submit Project", command=self.submit_project)
            submit_button.pack(side=tk.LEFT, padx=5)
            
            stop_button = ttk.Button(buttons_frame, text="Stop Assistant", command=self.stop_monitoring)
            stop_button.pack(side=tk.LEFT, padx=5)
            
            # Add drag and drop support
            self.setup_drag_and_drop(self.instructions_text)
            
            # === OUTPUT FILES TAB ===
            
            # Path frame
            path_frame = ttk.Frame(output_tab)
            path_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
            
            path_label = ttk.Label(path_frame, text="Output folder:")
            path_label.pack(side=tk.LEFT, padx=(0, 5))
            
            self.output_path_var = tk.StringVar()
            self.output_path_var.set(self.properties.get('outbox_folder', {}).get('default', 'Not set'))
            
            output_path_entry = ttk.Entry(path_frame, textvariable=self.output_path_var, width=50)
            output_path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
            
            browse_button = ttk.Button(path_frame, text="Browse", command=self.browse_output_folder)
            browse_button.pack(side=tk.LEFT, padx=(0, 5))
            
            refresh_button = ttk.Button(path_frame, text="Refresh", command=self.refresh_file_browser)
            refresh_button.pack(side=tk.LEFT)
            
            # File browser frame
            browser_frame = ttk.Frame(output_tab)
            browser_frame.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
            
            # Create a treeview for the file browser
            self.file_tree = ttk.Treeview(browser_frame, columns=("size", "modified"), selectmode="browse")
            self.file_tree.heading("#0", text="Name")
            self.file_tree.heading("size", text="Size")
            self.file_tree.heading("modified", text="Modified")
            
            self.file_tree.column("#0", width=300, minwidth=200)
            self.file_tree.column("size", width=100, minwidth=80, anchor=tk.E)
            self.file_tree.column("modified", width=150, minwidth=120)
            
            # Add scrollbars
            tree_yscroll = ttk.Scrollbar(browser_frame, orient="vertical", command=self.file_tree.yview)
            tree_xscroll = ttk.Scrollbar(browser_frame, orient="horizontal", command=self.file_tree.xview)
            self.file_tree.configure(yscrollcommand=tree_yscroll.set, xscrollcommand=tree_xscroll.set)
            
            # Grid layout for the file browser
            self.file_tree.grid(row=0, column=0, sticky="nsew")
            tree_yscroll.grid(row=0, column=1, sticky="ns")
            tree_xscroll.grid(row=1, column=0, sticky="ew")
            
            browser_frame.grid_rowconfigure(0, weight=1)
            browser_frame.grid_columnconfigure(0, weight=1)
            
            # Add right-click menu for file operations
            self.file_context_menu = tk.Menu(self.file_tree, tearoff=0)
            self.file_context_menu.add_command(label="Open", command=self.open_selected_file)
            self.file_context_menu.add_command(label="Open Folder", command=self.open_containing_folder)
            self.file_context_menu.add_separator()
            self.file_context_menu.add_command(label="Copy Path", command=self.copy_file_path)
            
            # Bind events
            self.file_tree.bind("<Double-1>", self.on_file_double_click)
            self.file_tree.bind("<Button-3>", self.on_file_right_click)
            
            # Action buttons
            action_frame = ttk.Frame(output_tab)
            action_frame.grid(row=2, column=0, sticky="ew", padx=5, pady=5)
            
            open_button = ttk.Button(action_frame, text="Open Selected", command=self.open_selected_file)
            open_button.pack(side=tk.LEFT, padx=5)
            
            open_folder_button = ttk.Button(action_frame, text="Open Containing Folder", command=self.open_containing_folder)
            open_folder_button.pack(side=tk.LEFT, padx=5)
            
            # === LOG TAB ===
            
            # Create and configure the output text widget
            self.output_text = scrolledtext.ScrolledText(log_tab, height=20, width=80, wrap=tk.WORD)
            self.output_text.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
            
            # Log tab buttons
            log_buttons_frame = ttk.Frame(log_tab)
            log_buttons_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=5)
            
            clear_log_button = ttk.Button(log_buttons_frame, text="Clear Log", command=self.clear_log)
            clear_log_button.pack(side=tk.LEFT, padx=5)
            
            # Initialize the temporary directory for project files
            self.temp_dir = Path(os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_projects"))
            self.temp_dir.mkdir(exist_ok=True)
            
            # Add initial message to log
            self.update_log("Assistant interface initialized and ready.")
            self.update_log(f"Monitoring inbox folder: {self.properties.get('inbox_folder', {}).get('default', 'Not set')}")
            self.update_log(f"Output folder: {self.properties.get('outbox_folder', {}).get('default', 'Not set')}")
            self.update_log("Use the 'Submit Project' tab to create a new project.")
            
            # Initialize the file browser
            self.refresh_file_browser()
            
    def setup_drag_and_drop(self, widget):
        """Setup drag and drop functionality for the given widget."""
        try:
            # This will only work on Windows
            widget.drop_target_register("DND_Files")
            widget.dnd_bind('<<Drop>>', self.handle_drop)
        except:
            # If drag and drop registration fails, just log it
            print("[DEBUG] Drag and drop registration failed. This feature may not be available on this platform.")
            
    def handle_drop(self, event):
        """Handle files dropped onto the widget."""
        try:
            # Get the dropped data
            data = event.data
            
            # Check if it's a file path or paths
            if data:
                # On Windows, paths might be enclosed in curly braces and separated by spaces
                if data.startswith('{') and data.endswith('}'):
                    data = data[1:-1]  # Remove the curly braces
                
                # Split by space, but handle paths with spaces correctly
                paths = []
                in_quotes = False
                current_path = ""
                
                for char in data:
                    if char == '"':
                        in_quotes = not in_quotes
                    elif char == ' ' and not in_quotes:
                        if current_path:
                            paths.append(current_path)
                            current_path = ""
                    else:
                        current_path += char
                
                if current_path:
                    paths.append(current_path)
                
                # Add each file to the project
                for path in paths:
                    path = path.strip('"')  # Remove any remaining quotes
                    self.add_file_to_project(path)
        except Exception as e:
            print(f"[DEBUG] Error handling dropped files: {str(e)}")
            self.update_log(f"Error handling dropped files: {str(e)}")

    def add_files(self):
        """Open a file dialog to add files to the project."""
        try:
            from tkinter import filedialog
            
            # Open file dialog
            files = filedialog.askopenfilenames(
                title="Select files to add to the project",
                filetypes=[
                    ("All Files", "*.*"),
                    ("Text Files", "*.txt"),
                    ("PDF Files", "*.pdf"),
                    ("Word Documents", "*.docx"),
                    ("Excel Files", "*.xlsx *.xls"),
                    ("CSV Files", "*.csv"),
                    ("Audio Files", "*.mp3 *.wav *.m4a *.ogg"),
                    ("Video Files", "*.mp4 *.avi *.mov *.mkv")
                ]
            )
            
            # Add each selected file to the project
            for file_path in files:
                self.add_file_to_project(file_path)
                
        except Exception as e:
            print(f"[DEBUG] Error adding files: {str(e)}")
            self.update_log(f"Error adding files: {str(e)}")

    def add_file_to_project(self, file_path):
        """Add a file to the project list."""
        try:
            file_path = Path(file_path)
            
            # Check if file exists
            if not file_path.exists() or not file_path.is_file():
                self.update_log(f"File not found: {file_path}")
                return
                
            # Check if file is already in the list
            if str(file_path) in self.project_files:
                self.update_log(f"File already added: {file_path.name}")
                return
                
            # Add to the list of project files
            self.project_files.append(str(file_path))
            
            # Add to the listbox
            self.files_listbox.insert(tk.END, f"{file_path.name} ({file_path})")
            
            self.update_log(f"Added file to project: {file_path.name}")
            
        except Exception as e:
            print(f"[DEBUG] Error adding file to project: {str(e)}")
            self.update_log(f"Error adding file to project: {str(e)}")

    def remove_selected_file(self):
        """Remove the selected file from the project."""
        try:
            # Get selected index
            selected = self.files_listbox.curselection()
            
            if not selected:
                self.update_log("No file selected to remove.")
                return
                
            # Get the index
            index = selected[0]
            
            # Remove from the project files list
            if 0 <= index < len(self.project_files):
                file_path = self.project_files.pop(index)
                
                # Remove from the listbox
                self.files_listbox.delete(index)
                
                self.update_log(f"Removed file from project: {Path(file_path).name}")
            
        except Exception as e:
            print(f"[DEBUG] Error removing file: {str(e)}")
            self.update_log(f"Error removing file: {str(e)}")

    def clear_files(self):
        """Clear all files from the project."""
        try:
            # Clear the project files list
            self.project_files = []
            
            # Clear the listbox
            self.files_listbox.delete(0, tk.END)
            
            self.update_log("Cleared all files from project.")
            
        except Exception as e:
            print(f"[DEBUG] Error clearing files: {str(e)}")
            self.update_log(f"Error clearing files: {str(e)}")

    def submit_project(self):
        """Submit the project for processing."""
        try:
            # Reset the current job folder and custom project name to ensure a new one is created
            self.current_job_folder = None
            self.custom_project_name = None
            
            # Get the instructions text
            instructions = self.instructions_text.get(1.0, tk.END).strip()
            
            if not instructions:
                self.update_log("Please enter instructions before submitting.")
                return
                
            # Create a timestamp for the project folder
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Generate a unique ID (first 8 chars of a UUID)
            import uuid
            unique_id = str(uuid.uuid4())[:8]
            
            # Get the inbox folder
            inbox_folder = self.properties.get('inbox_folder', {}).get('default', '')
            if not inbox_folder:
                self.update_log("No inbox folder configured. Please set the inbox folder in node properties.")
                return
                
            inbox_folder = Path(inbox_folder)
            if not inbox_folder.exists():
                self.update_log(f"Inbox folder does not exist: {inbox_folder}")
                return
                
            # Get custom project name if provided
            custom_name = self.project_name_var.get().strip()
            self.custom_project_name = custom_name if custom_name else None
            
            # Create a project folder in the inbox with custom name (if provided) plus timestamp and unique ID
            if custom_name:
                project_name = f"{custom_name}_{timestamp}_{unique_id}"
            else:
                project_name = f"Assistant_Project_{timestamp}_{unique_id}"
                
            project_folder = inbox_folder / project_name
            project_folder.mkdir(exist_ok=True)
            
            # Create the instructions file with file references
            instructions_file = project_folder / "Assistant_instructions.txt"
            
            with open(instructions_file, 'w', encoding='utf-8') as f:
                f.write(instructions)
                
                # Add file references if there are any
                if self.project_files:
                    f.write("\n\n")
                    for file_path in self.project_files:
                        file_name = Path(file_path).name
                        f.write(f"[{file_name}]\n")
            
            # Copy all project files to the project folder
            for file_path in self.project_files:
                source_path = Path(file_path)
                if source_path.exists() and source_path.is_file():
                    # Copy the file to the project folder
                    target_path = project_folder / source_path.name
                    shutil.copy2(str(source_path), str(target_path))
                    self.update_log(f"Copied file to project: {source_path.name}")
                else:
                    self.update_log(f"File not found, skipping: {source_path}")
            
            # Clear the form after submission
            self.instructions_text.delete(1.0, tk.END)
            self.project_name_var.set("")  # Clear the project name field
            self.clear_files()
            
            self.update_log(f"Project submitted successfully to: {project_folder}")
            self.update_log("Processing will begin shortly.")
            
        except Exception as e:
            print(f"[DEBUG] Error submitting project: {str(e)}")
            self.update_log(f"Error submitting project: {str(e)}")
            traceback.print_exc()

    def stop_monitoring(self):
        """Stop monitoring and close the window."""
        print("[DEBUG] Stopping assistant monitoring...")
        self.stop_event.set()
        
        # Get the workflow_id from inputs if available
        workflow_id = None
        if hasattr(self, 'inputs') and isinstance(self.inputs, dict) and 'workflow_id' in self.inputs:
            workflow_id = self.inputs.get('workflow_id')
            
        # If we have a workflow_id, mark it as stopped in the workflow manager
        if workflow_id:
            from main import workflow_manager
            try:
                workflow_manager.stop_workflow(workflow_id)
                print(f"[DEBUG] Marked workflow {workflow_id} as stopped in workflow manager")
            except Exception as e:
                print(f"[DEBUG] Error marking workflow as stopped: {str(e)}")
        
        if hasattr(self, 'monitor_window') and self.monitor_window:
            try:
                # Schedule the window destruction after a short delay
                # This allows any pending operations to complete
                self.monitor_window.after(100, self.destroy_monitor_window)
            except Exception as e:
                print(f"[DEBUG] Error scheduling window destruction: {str(e)}")
                self.destroy_monitor_window()

    def destroy_monitor_window(self):
        """Safely destroy the monitor window."""
        try:
            if hasattr(self, 'monitor_window') and self.monitor_window:
                print("[DEBUG] Destroying monitor window...")
                self.monitor_window.destroy()
                self.monitor_window = None
                print("[DEBUG] Monitor window destroyed successfully")
        except Exception as e:
            print(f"[DEBUG] Error destroying monitor window: {str(e)}")
            # Ensure the window reference is cleared even if destruction fails
            self.monitor_window = None

    def clear_log(self):
        """Clear the output text widget."""
        if hasattr(self, 'output_text'):
            self.output_text.delete(1.0, tk.END)

    def update_log(self, message):
        """Update the log with a new message."""
        if hasattr(self, 'output_text') and self.output_text:
            try:
                # Check if widget still exists and is mapped
                self.output_text.winfo_exists()
                self.monitor_window.winfo_exists()
                
                # Schedule the update on the main thread
                self.monitor_window.after(0, lambda: self._safe_update_text(message))
            except Exception as e:
                print(f"[DEBUG] Error updating log: {str(e)}")

    def _safe_update_text(self, message):
        """Safely update the text widget on the main thread."""
        try:
            self.output_text.insert(tk.END, f"{message}\n")
            self.output_text.see(tk.END)
        except Exception as e:
            print(f"[DEBUG] Error in _safe_update_text: {str(e)}")

    def is_audio_file(self, file_path):
        """Check if the file is a supported audio format."""
        SUPPORTED_AUDIO_EXTENSIONS = {'.mp3', '.mpeg', '.mpga', '.m4a', '.wav', '.webm'}
        return Path(file_path).suffix.lower() in SUPPORTED_AUDIO_EXTENSIONS

    def is_video_file(self, file_path):
        """Check if the file is a supported video format."""
        SUPPORTED_VIDEO_EXTENSIONS = {'.mp4', '.avi', '.mov', '.mkv', '.wmv', '.flv'}
        return Path(file_path).suffix.lower() in SUPPORTED_VIDEO_EXTENSIONS

    def is_csv_file(self, file_path):
        """Check if the file is a CSV file."""
        return Path(file_path).suffix.lower() == '.csv'

    def is_pdf_file(self, file_path):
        """Check if the file is a PDF file."""
        return Path(file_path).suffix.lower() == '.pdf'

    def is_excel_file(self, file_path):
        """Check if the file is an Excel file."""
        return Path(file_path).suffix.lower() in {'.xlsx', '.xls'}

    def is_docx_file(self, file_path):
        """Check if the file is a DOCX file."""
        return str(file_path).lower().endswith('.docx')

    def extract_audio_from_video(self, video_path, output_path):
        """Extract audio from a video file and save it as an MP3."""
        max_retries = 5
        base_delay = 1  # Start with 1 second delay
        
        for attempt in range(max_retries):
            try:
                print(f"[DEBUG] Attempt {attempt + 1}/{max_retries} to extract audio from video")
                video = AudioSegment.from_file(video_path)
                video.export(output_path, format="mp3")
                print(f"[DEBUG] Successfully extracted audio from video on attempt {attempt + 1}")
                return True
            except PermissionError as e:
                if attempt == max_retries - 1:  # Last attempt
                    print(f"[DEBUG] Error extracting audio from video: {str(e)}")
                    return False
                delay = base_delay * (2 ** attempt)  # Exponential backoff
                print(f"[DEBUG] Attempt {attempt + 1}/{max_retries}: Permission denied, waiting {delay}s before retry...")
                time.sleep(delay)
            except Exception as e:
                print(f"[DEBUG] Error extracting audio from video: {str(e)}")
                return False
        return False

    def wait_for_file_access(self, file_path, max_retries=5, initial_delay=1):
        """Wait for a file to become accessible with exponential backoff."""
        delay = initial_delay
        last_size = -1
        stable_count = 0
        
        for attempt in range(max_retries):
            try:
                current_size = Path(file_path).stat().st_size
                print(f"[DEBUG] Attempt {attempt + 1}/{max_retries}: File size = {current_size} bytes")
                
                # Check if file size has stabilized
                if current_size == last_size:
                    stable_count += 1
                    if stable_count >= 2:  # File size hasn't changed for 2 checks
                        print(f"[DEBUG] File size stable at {current_size} bytes")
                        # Try to open the file to verify access
                        with open(file_path, 'rb') as f:
                            return True
                else:
                    stable_count = 0
                
                last_size = current_size
                time.sleep(delay)
                delay *= 2  # Exponential backoff
                
            except (PermissionError, OSError) as e:
                print(f"[DEBUG] Attempt {attempt + 1}/{max_retries}: Access denied ({str(e)}), waiting {delay}s")
                if attempt == max_retries - 1:
                    print(f"[DEBUG] Max retries reached, giving up")
                    return False
                time.sleep(delay)
                delay *= 2  # Exponential backoff
                
        return False

    def process_audio_file(self, file_path, transcription_output_path):
        """Process an audio file using Whisper API and save it to the specified output path."""
        try:
            file_path = Path(file_path)
            project_folder = file_path.parent
            transcription_file = Path(transcription_output_path)
            self.update_log(f"Creating transcription file: {transcription_file.name}")
            print(f"[DEBUG] Transcription file will be: {transcription_file}")
            
            # Get the file size
            file_size = file_path.stat().st_size
            chunk_size = 25 * 1024 * 1024  # 25MB chunks
            
            if file_size > chunk_size:
                # Process large file in chunks
                self.update_log(f"Large file detected, processing in chunks: {file_path.name}")
                print(f"[DEBUG] Processing large file in chunks: {file_path}")
                audio = AudioSegment.from_file(str(file_path))
                chunk_duration = 10 * 60 * 1000  # 10 minutes in milliseconds
                
                # Process each chunk
                for i, chunk_start in enumerate(range(0, len(audio), chunk_duration)):
                    chunk_end = min(chunk_start + chunk_duration, len(audio))
                    chunk = audio[chunk_start:chunk_end]
                    
                    # Save chunk in project folder
                    chunk_file = project_folder / f"{file_path.stem}_chunk_{i}.mp3"
                    chunk.export(str(chunk_file), format="mp3")
                    
                    try:
                        # Process chunk
                        self.update_log(f"Processing chunk {i+1}: {chunk_file.name}")
                        print(f"[DEBUG] Processing chunk {i}: {chunk_file}")
                        transcript = self.transcribe_audio(str(chunk_file))
                        
                        # Append to transcription file
                        with open(transcription_file, 'a', encoding='utf-8') as f:
                            f.write(f"\n--- Chunk {i} ---\n")
                            f.write(transcript)
                        
                        self.update_log(f"Completed chunk {i+1}")
                        
                    finally:
                        # Clean up chunk file
                        if chunk_file.exists():
                            chunk_file.unlink()
                
            else:
                # Process small file directly
                self.update_log(f"Processing file: {file_path.name}")
                print(f"[DEBUG] Processing small file directly: {file_path}")
                transcript = self.transcribe_audio(str(file_path))
                with open(transcription_file, 'w', encoding='utf-8') as f:
                    f.write(transcript)
            
            # Process the complete transcription file
            self.update_log(f"Completed processing: {file_path.name}")
            print(f"[DEBUG] Completed transcription file generation: {transcription_file}")
            

            
            return True
            
        except Exception as e:
            self.update_log(f"Error in process_audio_file: {str(e)}")
            print(f"[DEBUG] Error in process_audio_file: {str(e)}")
            traceback.print_exc()
            raise

    def process_csv_file(self, file_path):
        """Process a CSV file by converting it to text."""
        try:
            file_path = Path(file_path)
            # Read CSV file
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to string representation
            text_content = f"CSV File Content from {file_path.name}:\n\n"
            text_content += "Column Names:\n" + ", ".join(str(col) for col in df.columns) + "\n\n"
            text_content += "Data:\n" + df.to_string(index=False) + "\n\n"
            text_content += f"Total Rows: {len(df)}\n"
            
            # Create text file in the same directory
            text_file_path = file_path.parent / f"{file_path.stem}.txt"
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Process the text file
            self.process_text_file(str(text_file_path))
            return True
            
        except Exception as e:
            self.update_log(f"Error processing CSV file: {str(e)}")
            print(f"[DEBUG] Error processing CSV file: {str(e)}")
            traceback.print_exc()
            raise

    def process_pdf_file(self, file_path):
        """Process a PDF file by converting it to text."""
        try:
            file_path = Path(file_path)
            text_content = f"PDF File Content from {file_path.name}:\n\n"
            
            # Read PDF file
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += f"--- Page {page_num + 1} ---\n"
                    text_content += page.extract_text() + "\n\n"
            
            # Create text file in the same directory
            text_file_path = file_path.parent / f"{file_path.stem}.txt"
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Process the text file
            self.process_text_file(str(text_file_path))
            return True
            
        except Exception as e:
            self.update_log(f"Error processing PDF file: {str(e)}")
            print(f"[DEBUG] Error processing PDF file: {str(e)}")
            traceback.print_exc()
            raise

    def process_excel_file(self, file_path):
        """Process an Excel file by converting it to text."""
        try:
            file_path = Path(file_path)
            text_content = f"Excel File Content from {file_path.name}:\n\n"
            
            # Read Excel file with context manager
            with pd.ExcelFile(file_path) as excel_file:
                # Process each sheet
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    text_content += f"Sheet: {sheet_name}\n"
                    text_content += "Column Names:\n" + ", ".join(str(col) for col in df.columns) + "\n\n"
                    text_content += "Data:\n" + df.to_string(index=False) + "\n\n"
                    text_content += f"Total Rows: {len(df)}\n\n"
                    text_content += "---\n\n"
                    # Explicitly delete DataFrame to free memory
                    del df
            
            # Create text file in the same directory
            text_file_path = file_path.parent / f"{file_path.stem}.txt"
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Process the text file
            self.process_text_file(str(text_file_path))
            return True
            
        except Exception as e:
            self.update_log(f"Error processing Excel file: {str(e)}")
            print(f"[DEBUG] Error processing Excel file: {str(e)}")
            traceback.print_exc()
            raise

    def process_text_file(self, file_path):
        """Process a text file."""
        try:
            file_path = Path(file_path)
            additional_content = []  # Initialize the list here
            
            # Wait for file to be fully written and accessible
            self.update_log(f"Waiting for file to be ready: {file_path.name}")
            if not self.wait_for_file_access(file_path, max_retries=10, initial_delay=2):
                raise Exception(f"File {file_path.name} is not accessible after waiting")
            
            # Read the file content with multiple encoding attempts
            self.update_log(f"Reading file: {file_path.name}")
            content = None
            encodings_to_try = ['utf-8', 'ascii', 'latin-1', 'cp1252']
            
            for encoding in encodings_to_try:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    print(f"[DEBUG] Successfully read file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    print(f"[DEBUG] Failed to read with {encoding} encoding, trying next...")
                    continue
                except Exception as e:
                    print(f"[DEBUG] Error reading with {encoding}: {str(e)}")
                    continue
            
            if content is None:
                raise Exception(f"Could not read file {file_path.name} with any supported encoding")
            
            # Preserve the raw user request exactly as provided in the input file
            original_user_request = content

            # Get and append pre-process text if available
            pre_process_text = self.properties.get('PreProcess', {}).get('default', '').strip()
            initial_response = None
            
            if pre_process_text:
                self.update_log("Appending pre-process text")
                content = f"{pre_process_text}\n\n{content}"

                # First round of API processing only if pre-process text exists
                api_endpoint = self.properties['text_api_endpoint']['default']
                if not api_endpoint:
                    raise ValueError("No text API endpoint configured")

                self.update_log("Processing initial content with text API")
                initial_response = self.send_to_api(content, api_endpoint)
                
                if not initial_response:
                    raise ValueError("No response received from initial API call")

                content = initial_response  # Use the API response for further processing
            
            # Process search tags and URLs only if we got an initial response
            if initial_response:
                # Get search API endpoint from properties
                search_api_endpoint = None
                for api_name, api_config in self.config.get('interfaces', {}).items():
                    if api_config.get('type') == 'SearchEngine':
                        search_api_endpoint = api_name
                        break
                
                if not search_api_endpoint:
                    print("[DEBUG] No search API endpoint configured")
                else:
                    # Extract search queries from response using regex
                    search_pattern = r'<search>(.*?)</search>'
                    search_queries = re.findall(search_pattern, initial_response)
                    
                    # Process each search query
                    for search_query in search_queries:
                        print(f"[DEBUG] Processing search query: {search_query}")
                        print(f"[DEBUG] Using search API endpoint: {search_api_endpoint}")
                        
                        # Get configured number of results
                        num_results = self.properties.get('num_search_results', {}).get('default', 3)
                        try:
                            num_results = int(num_results)
                        except (TypeError, ValueError):
                            num_results = 3
                            
                        # Send search query to search API
                        search_request = {
                            'content': search_query.strip(),
                            'num_results': num_results,
                            'skip': 0,
                            'format': 'json'
                        }
                        
                        print(f"[DEBUG] Sending search request to {search_api_endpoint}: {search_request}")
                        
                        # Send to API handler
                        search_results = process_api_request(
                            search_api_endpoint,
                            self.config,
                            search_request
                        )
                        
                        if search_results:
                            # Replace the search tag with the results
                            search_tag = f'<search>{search_query}</search>'
                            initial_response = initial_response.replace(search_tag, f'\n\nSearch Results:\n{search_results}\n\n')
                
                # Continue with URL processing only if we have an initial response
                urls = self.extract_urls(initial_response)
                
                if urls:
                    self.update_log(f"Found {len(urls)} URLs in processed content")
                    for url in urls:
                        if self.is_youtube_url(url):
                            self.update_log(f"Processing YouTube URL: {url}")
                            # Create YouTube transcript node
                            youtube_node = YoutubeTranscriptNode(
                                node_id="youtube_transcript_temp",
                                config=self.config
                            )
                            result = youtube_node.process({'input': url})
                            
                            youtube_content = []
                            if 'title' in result and result['title']:
                                youtube_content.append(f"Title: {result['title']}")
                            if 'description' in result and result['description']:
                                youtube_content.append(f"Description: {result['description']}")
                            if 'transcript' in result and not result['transcript'].startswith('Error:'):
                                youtube_content.append(f"Transcript: {result['transcript']}")
                            
                            if youtube_content:
                                content = f"\n\nYouTube Video Content from {url}:\n" + "\n\n".join(youtube_content)
                                additional_content.append(content)
                        else:
                            self.update_log(f"Processing web URL: {url}")
                            # Create web scraping node
                            web_node = WebScrapingNode(
                                node_id="web_scrape_temp",
                                config={
                                    'depth': 1,
                                    'max_retries': 2,
                                    'interfaces': self.config.get('interfaces', {})
                                }
                            )
                            web_node.get_depth_from_user = lambda: 1
                            result = web_node.process({'input': url})
                            
                            if result and isinstance(result, dict) and 'scraped_text' in result:
                                scraped_text = result['scraped_text'].strip()
                                if scraped_text:
                                    content = f"\n\nWeb Content from {url}:\n{scraped_text}"
                                    additional_content.append(content)
            
            # Combine content
            combined_content = [initial_response if initial_response else content]
            if additional_content:
                combined_content.append("\n=== Additional Content from URLs ===")
                combined_content.extend(additional_content)
            
            # Get the prompt template
            prompt = self.properties.get('Prompt', {}).get('default', '').strip()
            if prompt:
                combined_content.insert(0, prompt)
            
            # Join all content
            full_content = "\n\n".join(combined_content)
            self.update_log("Processing text content with API")
            print(f"[DEBUG] Sending content to API, length: {len(full_content)}")
            
            # Send to API
            api_endpoint = self.properties['text_api_endpoint']['default']
            if not api_endpoint:
                raise ValueError("No text API endpoint configured")
                
            print(f"[DEBUG] Using text API endpoint: {api_endpoint}")
            api_response = self.send_to_api(full_content, api_endpoint)
            
            if not api_response:
                raise ValueError("No response received from API")
            
            print(f"[DEBUG] Received API response of length: {len(api_response)}")
            
            # Process API output to extract any Excel code blocks and generate .xlsx files
            if hasattr(self, 'custom_project_name') and self.custom_project_name:
                base_name = self.custom_project_name
            else:
                base_name = file_path.stem

            filtered_md, excel_files = process_api_output(
                api_response,
                str(self.current_job_folder),
                base_name,
                insert_placeholders=False,
            )

            if excel_files:
                self.update_log(f"Generated {len(excel_files)} Excel file(s):")
                for p in excel_files:
                    try:
                        self.update_log(f"- {Path(p).name}")
                    except Exception:
                        pass

            # Create Word document only if there is meaningful content after filtering
            should_generate_docx = bool(filtered_md and filtered_md.strip())
            docx_filename = f"{base_name}.docx" if should_generate_docx else None
            temp_docx_path = (Path(self.current_job_folder) / docx_filename) if should_generate_docx else None
            if should_generate_docx:
                # Use ExportWord.py to create formatted document
                self.update_log("Creating formatted Word document using ExportWord")
                print(f"[DEBUG] Output path: {temp_docx_path}")
                convert_markdown_to_docx(filtered_md, str(temp_docx_path), formatting_enabled=True)
            else:
                self.update_log("No Word content detected; skipping Word document generation")
            
            # Move files to output folder if needed
            if self.current_job_folder:
                output_folder = Path(self.current_job_folder)
                
                # Move original text file if it's not a transcription
                if not file_path.name.endswith('_transcription.txt'):
                    # Instead of copying the original input .txt, write a combined
                    # log containing the exact API request and response so you can
                    # inspect both before downstream processing.
                    new_text_location = output_folder / file_path.name
                    try:
                        with open(new_text_location, 'w', encoding='utf-8') as f:
                            f.write("=== USER REQUEST ===\n")
                            f.write(original_user_request)
                            f.write("\n\n=== API OUTPUT ===\n")
                            f.write(api_response)
                    except Exception:
                        # Fallback: if writing combined file fails, at least copy the original
                        shutil.copy2(str(file_path), str(new_text_location))
                
                # Move Word document only if it was generated
                if should_generate_docx and temp_docx_path and temp_docx_path.exists():
                    new_docx_location = output_folder / docx_filename
                    shutil.move(str(temp_docx_path), str(new_docx_location))
                
                # Add this file to processed files list
                self.processed_files.append(file_path)
            
            self.update_log(f"Completed processing: {file_path.name}")
            
        except Exception as e:
            self.update_log(f"Error processing text file: {str(e)}")
            print(f"[DEBUG] Error processing text file: {str(e)}")
            print(f"[DEBUG] Error details: {str(e)}")
            traceback.print_exc()
            raise

    def extract_urls(self, text):
        """Extract URLs from text content."""
        # Regular expression for URLs with or without protocol
        url_pattern = r'\b(?:https?://|www\.)[\w\d\-_./#?=%&]+'
        
        urls_set = set()  # Use a set to store unique URLs
        matches = re.finditer(url_pattern, text, re.IGNORECASE)
        
        for match in matches:
            url = match.group(0)
            print(f"[DEBUG] Found potential URL: {url}")
            
            # Skip if it's just a domain without TLD or looks invalid
            if not re.match(r'^(?:https?:\/\/)?(?:[\w-]+\.)+[a-z]{2,}(?:\/[^\s]*)?$', url, re.IGNORECASE):
                print(f"[DEBUG] Skipping invalid URL: {url}")
                continue
                
            # Add https:// if no protocol specified
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'
                print(f"[DEBUG] Added https:// to URL: {url}")
            
            # Add to set to ensure uniqueness
            urls_set.add(url)
            
        urls = list(urls_set)  # Convert set back to list
        print(f"[DEBUG] Valid unique URLs found: {urls}")
        return urls

    def is_youtube_url(self, url):
        """Check if a URL is a YouTube URL."""
        # Add https:// if not present
        if not url.startswith(('http://', 'https://')):
            url = f'https://{url}'
            
        youtube_patterns = [
            r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/watch\?v=[\w-]+',
            r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/shorts\/[\w-]+',
            r'(?:https?:\/\/)?youtu\.be\/[\w-]+'
        ]
        
        return any(re.match(pattern, url) for pattern in youtube_patterns)

    def get_or_create_outbox_project_folder(self, inbox_project_folder):
        """Get or create the outbox project folder, maintaining consistent timestamps."""
        outbox_folder = Path(self.properties['outbox_folder']['default'])
        project_name = inbox_project_folder.name
        
        # Always add a new timestamp to ensure uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Generate a unique ID (first 8 chars of a UUID)
        import uuid
        unique_id = str(uuid.uuid4())[:8]
        
        # Extract base name (remove any existing timestamp if present)
        if '_202' in project_name:
            base_name = project_name.split('_202')[0]  # Get name without timestamp
        else:
            base_name = project_name
            
        # Create a new unique project folder name
        new_project_name = f"{base_name}_{timestamp}_{unique_id}"
        
        # Create new project folder
        project_folder = outbox_folder / new_project_name
        project_folder.mkdir(exist_ok=True, parents=True)
        print(f"[DEBUG] Created new unique output project folder: {project_folder}")
        
        return project_folder

    def create_output_structure(self, inbox_project_folder, file_path):
        """Create the output folder structure for a file."""
        file_path = Path(file_path)
        
        # Get or create project folder with consistent timestamp
        outbox_project_folder = self.get_or_create_outbox_project_folder(inbox_project_folder)
        
        # Create specific file subfolder using original file name without timestamp
        if '_202' in file_path.stem:
            base_name = file_path.stem.split('_202')[0]  # Remove any existing timestamp
        else:
            base_name = file_path.stem
            
        # Always add a new timestamp to ensure uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create a unique subfolder for this file
        file_subfolder = outbox_project_folder / f"{base_name}_{timestamp}"
        file_subfolder.mkdir(exist_ok=True, parents=True)
        print(f"[DEBUG] Created unique output subfolder: {file_subfolder}")
        
        return file_subfolder

    def create_job_folder(self, base_filename):
        """Create a timestamped job folder and return its path."""
        outbox_folder = self.properties['outbox_folder']['default']
        if not outbox_folder:
            raise Exception("Outbox folder not configured")

        # Create timestamp subfolder if not already created for this job
        if not self.current_job_folder:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Extract unique ID from project folder name if it exists
            project_folder_name = base_filename
            unique_id = ""
            
            # Try to extract the unique ID from the project folder name
            if "_" in project_folder_name:
                parts = project_folder_name.split("_")
                if len(parts) >= 4:  # Format: Assistant_Project_TIMESTAMP_UNIQUEID
                    unique_id = parts[-1]
                    
            # If no unique ID was found, generate a new one
            if not unique_id:
                import uuid
                unique_id = str(uuid.uuid4())[:8]
            
            # Use the same unique ID for the output folder to maintain the connection
            output_folder_name = f"Assistant_Output_{timestamp}_{unique_id}"
            self.current_job_folder = Path(outbox_folder) / output_folder_name
            self.current_job_folder.mkdir(exist_ok=True, parents=True)
            print(f"[DEBUG] Created job folder: {self.current_job_folder}")
            
        return self.current_job_folder

    def organize_inbox_file(self, file_path):
        """Create a project folder for a single file in inbox and move it there."""
        file_path = Path(file_path)
        if not file_path.exists():
            return None
            
        # If file is already in a project folder (not directly in inbox), return that folder
        inbox_folder = Path(self.properties['inbox_folder']['default'])
        relative_path = file_path.relative_to(inbox_folder)
        
        if len(relative_path.parts) > 1:
            # File is already in a subfolder
            project_folder = file_path.parent
            print(f"[DEBUG] File already in project folder: {project_folder}")
            return project_folder
            
        # First check if we can access the file before creating a project folder
        try:
            with open(file_path, 'rb') as f:
                # Just try to read 1 byte to verify access
                f.read(1)
        except (PermissionError, OSError) as e:
            print(f"[DEBUG] Process error: {str(e)}")
            return None
            
        # Create project folder for single file only after verifying access
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_folder = inbox_folder / f"{file_path.stem}_{timestamp}"
        project_folder.mkdir(exist_ok=True)
        
        # Move file to project folder
        try:
            new_file_path = project_folder / file_path.name
            shutil.move(str(file_path), str(new_file_path))
            print(f"[DEBUG] Moved file to new project folder: {new_file_path}")
            return project_folder
        except Exception as e:
            print(f"[DEBUG] Error moving file to project folder: {str(e)}")
            try:
                # Try to clean up the project folder if move failed
                if project_folder.exists():
                    project_folder.rmdir()
            except:
                pass
            return None

    def scan_inbox_folders(self):
        """Scan inbox and its subfolders for files to process."""
        try:
            inbox_folder = Path(self.properties['inbox_folder']['default'])

            # First, check for Assistant_instructions.txt directly in inbox
            instructions_file = inbox_folder / "Assistant_instructions.txt"
            if instructions_file.exists() and instructions_file not in self.processed_files:
                self.update_log(f"Found instructions file in inbox root: {instructions_file.name}")
                print(f"[DEBUG] Found instructions file in inbox root: {instructions_file}")
                
                referenced_files = self.parse_instructions_file(instructions_file)
                
                if referenced_files:
                    if not self.are_files_ready_for_processing(inbox_folder, referenced_files):
                        self.update_log(f"Referenced files for {instructions_file.name} are still being modified, waiting for sync to complete")
                        print(f"[DEBUG] Referenced files are still being modified, waiting for sync to complete")
                        return
                else:
                    self.update_log(f"No referenced files found in {instructions_file.name}, proceeding immediately")
                    print(f"[DEBUG] No referenced files found, proceeding immediately")
                
                project_folder = self.organize_instructions_with_files(instructions_file, referenced_files)
                
                if project_folder:
                    self.update_log(f"Created project folder for instructions file: {project_folder.name}")
                    print(f"[DEBUG] Created project folder for instructions file: {project_folder}")
                    self.process_project_folder(project_folder)
                return

            # Handle other individual files directly in inbox
            for item in inbox_folder.glob('*'):
                if item.is_file():
                    print(f"[DEBUG] Found file in inbox root: {item}")
                    project_folder = self.organize_inbox_file(item)
                    if project_folder:
                        print(f"[DEBUG] Processing organized project folder: {project_folder}")
                        self.process_project_folder(project_folder)

            # Then scan all subfolders
            print(f"[DEBUG] Scanning subfolders in inbox: {inbox_folder}")
            # Get a list of subdirectories. We use a list to avoid issues with the iterator changing size during processing.
            subfolders = [d for d in inbox_folder.iterdir() if d.is_dir()]

            for project_folder in subfolders:
                # The folder might have been deleted in a previous iteration by another process or this one
                if not project_folder.exists():
                    print(f"[DEBUG] Skipping already deleted folder: {project_folder}")
                    continue

                print(f"[DEBUG] Scanning project folder: {project_folder}")
                
                # Check if the folder is ready for processing (no recent file changes)
                if not self.is_folder_ready_for_processing(project_folder):
                    self.update_log(f"Project folder {project_folder.name} has recent changes, waiting for sync to complete")
                    print(f"[DEBUG] Project folder has recent changes, waiting for sync to complete")
                    continue
                
                # Process the folder. This might result in its deletion.
                self.process_project_folder(project_folder)

        except Exception as e:
            print(f"[DEBUG] Process error in scan_inbox_folders: {str(e)}")

    def process_project_folder(self, project_folder):
        """Process all files in a project folder."""
        try:
            project_folder = Path(project_folder)
            if not project_folder.exists():
                self.update_log(f"Project folder no longer exists: {project_folder}")
                return

            instructions_file = project_folder / "Assistant_instructions.txt"
            if instructions_file.exists():
                self.update_log(f"Found instructions file: {instructions_file.name}")
                self.process_instructions_file(instructions_file)
                return

            files_to_process = [f for f in project_folder.glob('*') if f.is_file()]
            total_files = len(files_to_process)
            if total_files == 0:
                return

            self.update_log(f"Found {total_files} files to process in folder: {project_folder}")
            processed_count = 0

            for file_path in files_to_process:
                if file_path in self.processed_files:
                    continue

                self.update_log(f"Processing file {processed_count + 1} of {total_files}: {file_path.name}")
                if not self.is_file_ready_for_processing(file_path):
                    self.update_log(f"File {file_path.name} is still being modified, waiting...")
                    continue

                try:
                    output_folder = self.create_output_structure(project_folder, file_path)
                    self.current_job_folder = output_folder

                    # Copy the original file to the output folder to preserve it
                    try:
                        shutil.copy(str(file_path), str(output_folder / file_path.name))
                        self.update_log(f"Copied original file to output folder: {file_path.name}")
                    except Exception as e:
                        self.update_log(f"Error copying original file: {e}")

                    if self.is_audio_file(file_path) or self.is_video_file(file_path):
                        self.update_log(f"Transcribing audio/video file: {file_path.name}")
                        transcription_output_path = output_folder / f"{file_path.stem}_transcription.txt"
                        self.process_audio_file(str(file_path), str(transcription_output_path))
                        self.update_log(f"Processing transcription for: {file_path.name}")
                        self.process_text_file(str(transcription_output_path))
                    elif self.is_csv_file(file_path):
                        self.update_log(f"Processing CSV file: {file_path.name}")
                        self.process_csv_file(str(file_path))
                    elif self.is_pdf_file(file_path):
                        self.update_log(f"Processing PDF file: {file_path.name}")
                        self.process_pdf_file(str(file_path))
                    elif self.is_excel_file(file_path):
                        self.update_log(f"Processing Excel file: {file_path.name}")
                        self.process_excel_file(str(file_path))
                    elif self.is_docx_file(file_path):
                        self.update_log(f"Processing DOCX file: {file_path.name}")
                        self.process_docx_file(str(file_path))
                    else:
                        self.update_log(f"Processing text file: {file_path.name}")
                        self.process_text_file(str(file_path))

                    processed_count += 1
                    self.update_log(f"Successfully processed: {file_path.name}")
                    self.processed_files.append(file_path)

                except Exception as e:
                    self.update_log(f"Error processing file {file_path.name}: {str(e)}")
                    print(f"[DEBUG] Error details: {e}")
                finally:
                    self.current_job_folder = None
                    self.custom_project_name = None  # Reset custom project name for next job

            if processed_count == total_files:
                self.update_log(f"All {total_files} files processed, deleting project folder: {project_folder}")
                self.delete_project_folder(project_folder)

        except Exception as e:
            self.update_log(f"Error in process_project_folder: {str(e)}")
            print(f"[DEBUG] Process error: {str(e)}")
        finally:
            # Always reset the current_job_folder and custom_project_name when done processing
            self.current_job_folder = None
            self.custom_project_name = None

    def process_instructions_file(self, instructions_file):
        """Process an instructions file that contains references to other files.
        
        The instructions file format is expected to be:
        - User content at the beginning
        - Followed by file references in the format [file_name1], [file_name2], etc.
        - May include a meta tag #PROJECT_NAME: <name> at the top
        """
        try:
            instructions_file = Path(instructions_file)
            project_folder = instructions_file.parent
            
            self.update_log(f"Processing instructions file: {instructions_file.name}")
            print(f"[DEBUG] Processing instructions file: {instructions_file}")
            
            # Read the instructions file
            with open(instructions_file, 'r', encoding='utf-8') as f:
                instructions_content = f.read()
            # Check for project name meta tag
            custom_name = None
            lines = instructions_content.splitlines()
            if lines and lines[0].startswith('#PROJECT_NAME:'):
                custom_name = lines[0][len('#PROJECT_NAME:'):].strip()
                self.custom_project_name = custom_name
                # Remove the meta tag from the instructions content for further processing
                instructions_content = '\n'.join(lines[1:])
            else:
                self.custom_project_name = None
            
            # Extract file references using regex
            file_references = re.findall(r'\[(.*?)\]', instructions_content)
            
            if file_references:
                self.update_log(f"Found {len(file_references)} file references in instructions file")
                print(f"[DEBUG] Found {len(file_references)} file references: {file_references}")
            else:
                print(f"[DEBUG] No file references found in instructions file")
            
            # Process each referenced file
            self.update_log(f"Found {len(file_references)} file references in instructions file")
            print(f"[DEBUG] Found {len(file_references)} file references: {file_references}")
            
            # Create a temporary directory for processed content
            temp_dir = project_folder / "temp_processed"
            temp_dir.mkdir(exist_ok=True)
            
            # Create output structure for the instructions file
            output_folder = self.create_output_structure(project_folder, instructions_file)
            self.current_job_folder = output_folder
            
            # Process each referenced file and collect their content
            file_contents = []
            missing_files = []
            
            for file_name in file_references:
                # Try to find the file in the project folder
                # First, try exact match
                file_path = project_folder / file_name
                
                # If not found, try case-insensitive search
                if not file_path.exists():
                    print(f"[DEBUG] File not found with exact name: {file_name}")
                    found = False
                    for item in project_folder.glob('*'):
                        if item.is_file() and item.name.lower() == file_name.lower():
                            file_path = item
                            found = True
                            print(f"[DEBUG] Found file with case-insensitive match: {item.name}")
                            break
                    # If still not found, try fuzzy matching
                    if not found:
                        print(f"[DEBUG] Trying fuzzy matching for: {file_name}")
                        best_match = None
                        best_ratio = 0
                        for item in project_folder.glob('*'):
                            if item.is_file():
                                ratio = difflib.SequenceMatcher(None, file_name.lower(), item.name.lower()).ratio()
                                if ratio > 0.8 and ratio > best_ratio:  # 80% similarity threshold
                                    best_match = item
                                    best_ratio = ratio
                        
                        if best_match:
                            file_path = best_match
                            found = True
                            print(f"[DEBUG] Found file with fuzzy match ({best_ratio:.2f}): {best_match.name}")
                
                if not file_path.exists():
                    self.update_log(f"Referenced file not found: {file_name}")
                    print(f"[DEBUG] Referenced file not found: {file_name}")
                    missing_files.append(file_name)
                    continue

                # Copy the original file to the output folder to preserve it
                try:
                    shutil.copy(str(file_path), str(output_folder / file_path.name))
                    self.update_log(f"Copied original file to output folder: {file_path.name}")
                except Exception as e:
                    self.update_log(f"Error copying original file: {e}")
            
                self.update_log(f"Processing referenced file: {file_path.name}")
                print(f"[DEBUG] Processing referenced file: {file_path}")
                
                # Process the file based on type and save content to a text file
                content_file = temp_dir / f"{file_path.stem}_content.txt"
                
                try:
                    if self.is_audio_file(file_path) or self.is_video_file(file_path):
                        # For audio/video, transcribe and save to text file
                        self.update_log(f"Processing audio/video file: {file_path.name}")
                        if self.is_video_file(file_path):
                            # Create extracted audio file path
                            audio_file = temp_dir / f"{file_path.stem}_extracted_audio.mp3"
                            print(f"[DEBUG] Extracting audio to: {audio_file}")
                            
                            # Extract audio from video
                            if self.extract_audio_from_video(file_path, audio_file):
                                self.update_log(f"Audio extraction complete, transcribing: {audio_file.name}")
                                # Use process_audio_file to handle chunking and write directly to the content_file
                                self.process_audio_file(str(audio_file), str(content_file))
                                
                                # Read the generated transcription file to get the content
                                if content_file.exists():
                                    with open(content_file, 'r', encoding='utf-8') as tf:
                                        transcript = tf.read()
                                else:
                                    transcript = "[ERROR] Transcription file not found after processing."
                                
                                # Delete the extracted audio file after processing
                                if audio_file.exists():
                                    audio_file.unlink()
                                    print(f"[DEBUG] Deleted extracted audio file: {audio_file}")
                                
                                # Delete the original video file after processing
                                if file_path.exists():
                                    file_path.unlink()
                                    print(f"[DEBUG] Deleted original video file: {file_path}")
                            else:
                                self.update_log(f"Failed to extract audio from video: {file_path.name}")
                                with open(content_file, 'w', encoding='utf-8') as f:
                                    f.write(f"Error: Failed to extract audio from video file: {file_path.name}")
                                continue
                        else:
                            # Use process_audio_file to handle chunking and write directly to the content_file
                            self.process_audio_file(str(file_path), str(content_file))
                        
                        with open(content_file, 'r', encoding='utf-8') as f:
                            transcript = f.read()
                        
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(f"Content from {file_path.name}:\n\n{transcript}")
                    elif self.is_csv_file(file_path):
                        # For CSV, convert to text
                        df = pd.read_csv(file_path)
                        text_content = f"CSV File Content from {file_path.name}:\n\n"
                        text_content += "Column Names:\n" + ", ".join(str(col) for col in df.columns) + "\n\n"
                        text_content += "Data:\n" + df.to_string(index=False) + "\n\n"
                        text_content += f"Total Rows: {len(df)}\n"
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                    elif self.is_pdf_file(file_path):
                        # For PDF, extract text
                        text_content = f"PDF File Content from {file_path.name}:\n\n"
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            
                            # Extract text from each page
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                text_content += f"--- Page {page_num + 1} ---\n"
                                text_content += page.extract_text() + "\n\n"
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                    elif self.is_excel_file(file_path):
                        # For Excel, convert to text
                        text_content = f"Excel File Content from {file_path.name}:\n\n"
                        with pd.ExcelFile(file_path) as excel_file:
                            # Process each sheet
                            for sheet_name in excel_file.sheet_names:
                                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                                text_content += f"Sheet: {sheet_name}\n"
                                text_content += "Column Names:\n" + ", ".join(str(col) for col in df.columns) + "\n\n"
                                text_content += "Data:\n" + df.to_string(index=False) + "\n\n"
                                text_content += f"Total Rows: {len(df)}\n\n"
                                text_content += "---\n\n"
                                # Explicitly delete DataFrame to free memory
                                del df
                        
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                    elif self.is_docx_file(file_path):
                        # For DOCX, extract text
                        doc = Document(file_path)
                        text_content = f"DOCX File Content from {file_path.name}:\n\n"
                        for para in doc.paragraphs:
                            text_content += para.text + "\n\n"
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                    else:
                        # For text files, just read the content
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text_content = f"Text File Content from {file_path.name}:\n\n{f.read()}"
                        with open(content_file, 'w', encoding='utf-8') as f:
                            f.write(text_content)
                    
                    # Add to the list of file contents
                    with open(content_file, 'r', encoding='utf-8') as f:
                        file_contents.append(f.read())
                    
                    # No longer copy the original file to the output folder
                    # Removed: shutil.copy2(str(file_path), str(output_folder / file_path.name))
                    
                except Exception as e:
                    self.update_log(f"Error processing referenced file {file_path.name}: {str(e)}")
                    print(f"[DEBUG] Error processing referenced file: {str(e)}")
                    traceback.print_exc()
            
            # Create the combined content file
            combined_file = project_folder / f"{instructions_file.stem}_combined.txt"
            
            with open(combined_file, 'w', encoding='utf-8') as f:
                # Write the original instructions content
                f.write(instructions_content)
                
                # Add a separator
                f.write("\n\n--- REFERENCED FILE CONTENTS ---\n\n")
                
                # Add the content of each processed file
                for content in file_contents:
                    f.write(content)
                    f.write("\n\n--- END OF FILE ---\n\n")
                
                # Add information about missing files if any
                if missing_files:
                    f.write("\n\n--- MISSING FILES ---\n\n")
                    for file_name in missing_files:
                        f.write(f"File not found: {file_name}\n")
            
            # Process the combined file
            self.update_log(f"Processing combined content file")
            print(f"[DEBUG] Processing combined content file: {combined_file}")
            self.process_text_file(str(combined_file))
            
            # Copy only the combined file to the output folder
            # Important: process_text_file() already wrote a combined API request+response
            # file named "<stem>_combined.txt" into output_folder. To avoid overwriting it,
            # store the pre-API combined source as a separate file.
            combined_output_file = output_folder / f"{instructions_file.stem}_combined_source.txt"
            try:
                if not combined_output_file.exists():
                    shutil.copy2(str(combined_file), str(combined_output_file))
                    print(f"[DEBUG] Copied source combined file to output folder: {combined_output_file}")
            except Exception as e:
                print(f"[DEBUG] Error copying source combined file: {e}")
            
            # No longer copy the original instructions file to the output folder
            # Removed: shutil.copy2(str(instructions_file), str(output_folder / instructions_file.name))
            
            # Clean up temporary directory
            try:
                shutil.rmtree(str(temp_dir))
            except Exception as e:
                print(f"[DEBUG] Error cleaning up temp directory: {str(e)}")
            
            # Delete the project folder after processing
            self.delete_project_folder(project_folder)
            
        except Exception as e:
            self.update_log(f"Error processing instructions file: {str(e)}")
            print(f"[DEBUG] Error processing instructions file: {str(e)}")
            traceback.print_exc()
        finally:
            self.current_job_folder = None
            self.custom_project_name = None  # Reset custom project name for next job

    def parse_instructions_file(self, instructions_file):
        """Parse an instructions file to find referenced files.
        
        Args:
            instructions_file: Path to the instructions file
            
        Returns:
            List of referenced filenames
        """
        try:
            # Read the instructions file
            with open(instructions_file, 'r', encoding='utf-8') as f:
                instructions_content = f.read()
            
            # Extract file references using regex
            file_references = re.findall(r'\[(.*?)\]', instructions_content)
            
            if file_references:
                self.update_log(f"Found {len(file_references)} file references in instructions file")
                print(f"[DEBUG] Found {len(file_references)} file references: {file_references}")
            else:
                print(f"[DEBUG] No file references found in instructions file")
            
            return file_references
            
        except Exception as e:
            self.update_log(f"Error parsing instructions file: {str(e)}")
            print(f"[DEBUG] Error parsing instructions file: {str(e)}")
            traceback.print_exc()
            return []

    def organize_instructions_with_files(self, instructions_file, referenced_files):
        """Create a project folder for instructions file and all referenced files.
        
        Args:
            instructions_file: Path to the instructions file
            referenced_files: List of filenames referenced in the instructions file
            
        Returns:
            Path to the created project folder
        """
        try:
            instructions_file = Path(instructions_file)
            inbox_folder = Path(self.properties['inbox_folder']['default'])
            
            # Create a timestamp for the project folder
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Create project folder name based on instructions file name
            project_folder_name = f"{instructions_file.stem}_{timestamp}"
            project_folder = inbox_folder / project_folder_name
            
            # Create the project folder
            project_folder.mkdir(exist_ok=True)
            
            # First, find all referenced files in the inbox folder
            found_files = []
            missing_files = []
            
            for file_name in referenced_files:
                # Try to find the file in the inbox folder
                # First, try exact match
                file_path = inbox_folder / file_name
                
                # If not found, try case-insensitive search
                if not file_path.exists():
                    print(f"[DEBUG] File not found with exact name: {file_name}")
                    found = False
                    for item in inbox_folder.glob('*'):
                        if item.is_file() and item.name.lower() == file_name.lower():
                            file_path = item
                            found = True
                            print(f"[DEBUG] Found file with case-insensitive match: {item.name}")
                            break
                    # If still not found, try fuzzy matching
                    if not found:
                        print(f"[DEBUG] Trying fuzzy matching for: {file_name}")
                        best_match = None
                        best_ratio = 0
                        for item in inbox_folder.glob('*'):
                            if item.is_file():
                                ratio = difflib.SequenceMatcher(None, file_name.lower(), item.name.lower()).ratio()
                                if ratio > 0.8 and ratio > best_ratio:  # 80% similarity threshold
                                    best_match = item
                                    best_ratio = ratio
                        
                        if best_match:
                            file_path = best_match
                            found = True
                            print(f"[DEBUG] Found file with fuzzy match ({best_ratio:.2f}): {best_match.name}")
                
                if file_path.exists():
                    found_files.append(file_path)
                    print(f"[DEBUG] Found referenced file: {file_path}")
                else:
                    missing_files.append(file_name)
                    print(f"[DEBUG] Referenced file not found: {file_name}")
            
            # Move the instructions file to the project folder
            dest_instructions_file = project_folder / instructions_file.name
            shutil.copy2(str(instructions_file), str(dest_instructions_file))
            os.remove(str(instructions_file))
            print(f"[DEBUG] Moved instructions file to project folder: {dest_instructions_file}")
            
            # Move all found files to the project folder
            for file_path in found_files:
                dest_file = project_folder / file_path.name
                shutil.copy2(str(file_path), str(dest_file))
                os.remove(str(file_path))
                print(f"[DEBUG] Moved referenced file to project folder: {dest_file}")
            
            # Create a note about missing files if any
            if missing_files:
                missing_note_path = project_folder / "missing_files.txt"
                with open(missing_note_path, 'w', encoding='utf-8') as f:
                    f.write("The following referenced files were not found in the inbox:\n\n")
                    for file_name in missing_files:
                        f.write(f"- {file_name}\n")
                print(f"[DEBUG] Created note about missing files: {missing_note_path}")
            
            # Add the project folder to processed_files to avoid reprocessing
            self.processed_files.append(instructions_file)
            
            return project_folder
            
        except Exception as e:
            self.update_log(f"Error organizing instructions with files: {str(e)}")
            print(f"[DEBUG] Error organizing instructions with files: {str(e)}")
            traceback.print_exc()
            return None

    def transcribe_audio(self, file_path):
        """Transcribe audio using the configured Whisper API endpoint."""
        try:
            # Get the selected whisper endpoint from node properties
            whisper_prop = self.properties.get('whisper_api_endpoint', {})
            whisper_endpoint_name = whisper_prop.get('value') or whisper_prop.get('default')
            if not whisper_endpoint_name:
                raise ValueError("Whisper API endpoint is not configured in node properties.")

            # Get the configuration for the selected endpoint
            api_config = self.config['interfaces'].get(whisper_endpoint_name)
            if not api_config:
                raise ValueError(f"Configuration for endpoint '{whisper_endpoint_name}' not found.")

            # Get the selected model for the endpoint
            model = api_config.get('selected_model')
            if not model:
                raise ValueError(f"No model selected for API endpoint '{whisper_endpoint_name}'.")

            self.update_log(f"Transcribing with model '{model}' via '{whisper_endpoint_name}'...")

            # Use the modern, inherited send_api_request method
            # Pass the file_path as a keyword argument so it's handled as part of `additional_params`
            response = super().send_api_request(
                content=None,  # No text content needed for transcription
                api_name=whisper_endpoint_name,
                model=model,
                file=file_path  # Pass file path here
            )

            if response.success:
                # The response content for Whisper is the transcribed text
                return response.content
            else:
                error_message = f"Whisper API Error: {response.error}"
                self.update_log(error_message)
                print(f"[AssistantNode] {error_message}")
                raise Exception(error_message)

        except Exception as e:
            error_message = f"An unexpected error occurred in transcribe_audio: {str(e)}"
            self.update_log(error_message)
            print(f"[AssistantNode] {error_message}")
            traceback.print_exc()
            raise

    class FileHandler(FileSystemEventHandler):
        def __init__(self, node):
            super().__init__()
            self.node = node
            self.processing_files = set()  # Track files being processed

        def _handle_file(self, file_path):
            """Common handler for file events"""
            try:
                file_path = str(Path(file_path))  # Normalize path
                
                # Skip if currently being processed or already processed
                if file_path in self.processing_files or file_path in self.node.processed_files:
                    return
                
                if not Path(file_path).exists():
                    return
                
                # Skip temporary files
                if file_path.endswith('.tmp') or file_path.endswith('.crdownload'):
                    return
                
                # Add to processing set
                self.processing_files.add(file_path)
                
                try:
                    # If file is in inbox root, organize it first
                    inbox_folder = Path(self.node.properties['inbox_folder']['default'])
                    file_path_obj = Path(file_path)
                    
                    if file_path_obj.parent == inbox_folder:
                        print(f"[DEBUG] Organizing new file in inbox: {file_path}")
                        project_folder = self.node.organize_inbox_file(file_path)
                        if project_folder:
                            # File has been moved to a project folder, it will be picked up by the next event
                            return
                    else:
                        # File is in a project folder, process it
                        print(f"[DEBUG] Processing file in project folder: {file_path}")
                        output_folder = self.node.create_output_structure(file_path_obj.parent, file_path_obj)
                        self.node.current_job_folder = output_folder

                        # Copy the original file to the output folder to preserve it
                        try:
                            shutil.copy(str(file_path), str(output_folder / file_path_obj.name))
                            self.node.update_log(f"Copied original file to output folder: {file_path_obj.name}")
                        except Exception as e:
                            self.node.update_log(f"Error copying original file: {e}")
                        
                        try:
                            if self.node.is_audio_file(file_path):
                                self.node.update_log(f"Processing audio file: {file_path}")
                                # Create a temporary file for the transcription
                                transcription_output_path = output_folder / f"{Path(file_path).stem}_transcription.txt"
                                self.node.process_audio_file(file_path, transcription_output_path)
                                # Now process the transcription to generate the final output
                                self.node.process_text_file(str(transcription_output_path))
                                # Move or delete the original audio file to prevent reprocessing
                                try:
                                    file_path_obj = Path(file_path)
                                    # Delete the original file since we already copied it to the output folder
                                    if file_path_obj.exists():
                                        file_path_obj.unlink()
                                        self.node.update_log(f"Deleted original audio file after processing: {file_path_obj.name}")
                                except Exception as e:
                                    self.node.update_log(f"Warning: Could not delete original audio file: {str(e)}")
                            elif self.node.is_video_file(file_path):
                                self.node.update_log(f"Processing video file: {file_path}")
                                self.node.process_video_file(file_path)
                            elif self.node.is_csv_file(file_path):
                                self.node.update_log(f"Processing CSV file: {file_path}")
                                self.node.process_csv_file(file_path)
                            elif self.node.is_pdf_file(file_path):
                                self.node.update_log(f"Processing PDF file: {file_path}")
                                self.node.process_pdf_file(file_path)
                            elif self.node.is_excel_file(file_path):
                                self.node.update_log(f"Processing Excel file: {file_path}")
                                self.node.process_excel_file(file_path)
                            elif self.node.is_docx_file(file_path):
                                self.node.update_log(f"Processing DOCX file: {file_path}")
                                self.node.process_docx_file(file_path)
                            else:
                                self.node.update_log(f"Processing text file: {file_path}")
                                self.node.process_text_file(file_path)
                                
                        finally:
                            self.node.current_job_folder = None
                            
                    # Add to processed files only after successful processing
                    if file_path not in self.node.processed_files:
                        self.node.processed_files.append(file_path)
                        self.node.update_log(f"Successfully processed: {file_path}")
                    
                except Exception as e:
                    self.node.update_log(f"Error processing file {file_path}: {str(e)}")
                    print(f"[DEBUG] Error details: {e}")
            
            except Exception as e:
                self.node.update_log(f"Error in file handler: {str(e)}")
                print(f"[DEBUG] Handler error details: {e}")
            
            finally:
                # Remove from processing set whether successful or not
                self.processing_files.discard(file_path)

        def on_created(self, event):
            if event.is_directory:
                return
            self._handle_file(event.src_path)

        def on_modified(self, event):
            if event.is_directory:
                return
            # Only handle modification events for files in project folders
            if "_transcription.txt" not in event.src_path:
                self._handle_file(event.src_path)

        def on_moved(self, event):
            if event.is_directory:
                return
            self._handle_file(event.dest_path)

    def process(self, inputs):
        """Start monitoring the input folder for text and audio files."""
        try:
            # Store the inputs dictionary for later access
            self.inputs = inputs
            
            # Get the workflow's stop_event if available
            workflow_stop_event = None
            if 'stop_event' in inputs:
                workflow_stop_event = inputs['stop_event']
                print("[DEBUG] Using workflow stop_event for monitoring")
            
            # Get the inbox folder path from properties
            inbox_folder = self.properties['inbox_folder']['default']
            if not inbox_folder:
                raise Exception("Inbox folder not configured")

            # Create the monitor window
            self.create_monitor_window()

            # Start scanning for existing files
            while not self.stop_event.is_set():
                # Also check the workflow's stop_event if available
                if workflow_stop_event and workflow_stop_event.is_set():
                    print("[DEBUG] Workflow stop_event detected, stopping assistant monitoring")
                    self.stop_event.set()  # Set our internal stop event as well
                    break
                    
                try:
                    self.scan_inbox_folders()
                    time.sleep(1)  # Wait before next scan
                except Exception as e:
                    print(f"[DEBUG] Error in scan cycle: {str(e)}")
                    # Don't re-raise, just continue monitoring
                    continue

            return {'output': 'Assistant monitoring stopped normally'}
            
        except Exception as e:
            print(f"[DEBUG] Critical error in process: {str(e)}")
            return {'output': f'Assistant error: {str(e)}'}

    def send_to_api(self, prompt, api_endpoint):
        """Send the prompt to the API and return the response."""
        try:
            api_config = self.config['interfaces'].get(api_endpoint, {})
            model = api_config.get('selected_model')
            max_tokens = api_config.get('max_tokens', 2000)
            temperature = api_config.get('temperature', 0.7)
            system_message = "You are a helpful assistant that preserves all markdown formatting in your responses. Always maintain the original markdown syntax including headers, lists, code blocks, tables, and other formatting. Never modify or strip markdown formatting."

            # Combine system message with the user's prompt
            full_prompt = f"{system_message}\n\nUser Prompt:\n{prompt}"

            # Use the modern, inherited send_api_request method
            response = super().send_api_request(
                content=full_prompt,
                api_name=api_endpoint,
                model=model,
                max_tokens=max_tokens,
                temperature=temperature
            )

            if response.success:
                return response.content
            else:
                self.update_log(f"API Error: {response.error}")
                print(f"[AssistantNode] API Error: {response.error}")
                return f"Error: {response.error}"

        except Exception as e:
            self.update_log(f"An unexpected error occurred in send_to_api: {str(e)}")
            print(f"[AssistantNode] An unexpected error occurred in send_to_api: {str(e)}")
            traceback.print_exc()
            return f"Error: {str(e)}"

    def update_property(self, prop_name: str, value: Any) -> None:
        """Update a property value and ensure it is properly saved."""
        # Handle property migrations
        if prop_name == 'search_url':
            print(f"[DEBUG] Ignoring deprecated property: {prop_name}")
            return
            
        if prop_name in self.properties:
            self.properties[prop_name]['default'] = value
            print(f"[DEBUG] Updated property {prop_name} to value: {value}")

    def get_properties(self) -> Dict[str, Any]:
        """Get the current properties, handling any deprecated ones."""
        props = super().get_properties()
        # Remove deprecated properties
        if 'search_url' in props:
            del props['search_url']
        return props

    def setup_token_log(self):
        """
        Sets up the token usage logging directory structure.
        Creates a Logs directory in nodes folder and a subdirectory for this node.
        Returns the path to the token log CSV file.
        """
        try:
            # Get the node name from properties
            node_name = self.properties.get('node_name', {}).get('default', 'AssistantNode')
            
            # Create the Logs directory in the nodes folder
            nodes_dir = Path(__file__).parent
            logs_dir = nodes_dir / "Logs"
            logs_dir.mkdir(exist_ok=True)
            
            # Create node-specific subdirectory
            node_logs_dir = logs_dir / node_name
            node_logs_dir.mkdir(exist_ok=True)
            
            # Define the CSV file path
            log_file = node_logs_dir / "token_usage.csv"
            
            # If the file doesn't exist, create it with headers
            if not log_file.exists():
                with open(log_file, 'w', newline='', encoding='utf-8') as f:
                    import csv
                    writer = csv.writer(f)
                    writer.writerow(['ID', 'Date', 'Time', 'API_Endpoint', 'Model', 'SubmitTokens', 'ReplyTokens', 'TotalTokens', 'AudioDuration(s)'])
                print(f"[DEBUG] Created new token usage log file: {log_file}")
            
            return log_file
        except Exception as e:
            print(f"[DEBUG] Error setting up token log: {str(e)}")
            traceback.print_exc()
            return None

    def log_token_usage(self, api_endpoint, model, token_usage):
        """
        Logs token usage information to a CSV file.
        
        Args:
            api_endpoint: The API endpoint used
            model: The model used
            token_usage: Dictionary containing token usage information
        """
        try:
            log_file = self.setup_token_log()
            if not log_file:
                return
            
            # Extract token information
            prompt_tokens = token_usage.get('prompt_tokens', 0)
            completion_tokens = token_usage.get('completion_tokens', 0)
            total_tokens = token_usage.get('total_tokens', 0)
            audio_duration = token_usage.get('audio_duration', 0)
            
            # Get current time
            now = datetime.now()
            date_str = now.strftime("%Y-%m-%d")
            time_str = now.strftime("%H:%M:%S")
            
            # Generate a unique ID
            import uuid
            unique_id = str(uuid.uuid4())[:8]
            
            # Log the information
            with open(log_file, 'a', newline='', encoding='utf-8') as f:
                import csv
                writer = csv.writer(f)
                writer.writerow([
                    unique_id, 
                    date_str, 
                    time_str, 
                    api_endpoint, 
                    model, 
                    prompt_tokens, 
                    completion_tokens, 
                    total_tokens,
                    audio_duration
                ])
            
            # Add audio duration to log message if available
            if audio_duration > 0:
                minutes = audio_duration / 60
                self.update_log(f"Token usage logged: {prompt_tokens} input, {completion_tokens} output, {total_tokens} total, {audio_duration:.1f}s ({minutes:.2f}min) audio")
                print(f"[DEBUG] Token usage logged to {log_file} with {minutes:.2f} minutes of audio")
            else:
                self.update_log(f"Token usage logged: {prompt_tokens} input, {completion_tokens} output, {total_tokens} total")
                print(f"[DEBUG] Token usage logged to {log_file}")
        except Exception as e:
            print(f"[DEBUG] Error logging token usage: {str(e)}")
            traceback.print_exc()

    def extract_search_tags(self, text):
        """Extract and process search tags from text."""
        search_tags = re.findall(r'<search>(.*?)</search>', text)
        urls = []
        
        if search_tags:
            for search_query in search_tags:
                print(f"[DEBUG] Processing search query: {search_query}")
                search_api = self.properties.get('search_api_endpoint')
                if search_api:
                    print(f"[DEBUG] Using search API endpoint: {search_api}")
                    num_results = self.properties.get('num_search_results', 3)
                    search_request = {
                        'content': search_query.strip(),
                        'num_results': num_results,
                        'skip': 0,
                        'format': 'json'
                    }
                    print(f"[DEBUG] Sending search request to {search_api}: {search_request}")
                    search_response = process_api_request(search_api, search_request, self.config)
                    if isinstance(search_response, dict) and 'urls' in search_response:
                        urls.extend(search_response['urls'])
                        print(f"[DEBUG] Found {len(search_response['urls'])} valid URLs")
                    else:
                        print("[DEBUG] No valid URLs found in search response")
        
        print(f"[DEBUG] Valid URLs found: {urls}")
        return urls

    def process_docx_file(self, file_path):
        """Process a DOCX file."""
        try:
            file_path = Path(file_path)
            # Read DOCX file
            doc = Document(file_path)
            
            # Extract text from each paragraph
            text_content = f"DOCX File Content from {file_path.name}:\n\n"
            for para in doc.paragraphs:
                text_content += para.text + "\n\n"
            
            # Create text file in the same directory
            text_file_path = file_path.parent / f"{file_path.stem}.txt"
            with open(text_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Process the text file
            self.process_text_file(str(text_file_path))
            return True
            
        except Exception as e:
            self.update_log(f"Error processing DOCX file: {str(e)}")
            print(f"[DEBUG] Error processing DOCX file: {str(e)}")
            traceback.print_exc()
            raise

    def delete_project_folder(self, project_folder):
        """Delete a project folder after processing.
        
        Args:
            project_folder: Path to the project folder to delete
        """
        try:
            project_folder = Path(project_folder)
            inbox_folder = Path(self.properties['inbox_folder']['default'])
            
            # Safety checks before deletion
            if not project_folder.exists():
                print(f"[DEBUG] Project folder already deleted: {project_folder}")
                return True
                
            if project_folder == inbox_folder:
                print(f"[DEBUG] Cannot delete inbox folder: {project_folder}")
                return False
                
            if not project_folder.is_relative_to(inbox_folder):
                print(f"[DEBUG] Project folder not in inbox: {project_folder}")
                return False
            
            print(f"[DEBUG] Attempting to delete project folder: {project_folder}")
            
            # Clean up any references before deletion
            self.current_job_folder = None
            
            # Remove any processed files references from this folder
            self.processed_files = [f for f in self.processed_files if not str(f).startswith(str(project_folder))]
            
            # Add a small delay before attempting deletion to allow pending operations to complete
            print(f"[DEBUG] Waiting 2 seconds before deletion attempt...")
            time.sleep(2)
            
            # Delete the folder
            shutil.rmtree(str(project_folder))
            print(f"[DEBUG] Successfully deleted project folder: {project_folder}")
            return True
            
        except Exception as e:
            print(f"[DEBUG] Error deleting project folder: {str(e)}")
            traceback.print_exc()
            return False

    def is_file_ready_for_processing(self, file_path):
        """Check if a file is ready for processing by verifying it's not being modified.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            bool: True if the file is ready for processing, False otherwise
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return False
                
            # Get the current file size and modification time
            current_size = file_path.stat().st_size
            current_mtime = file_path.stat().st_mtime
            
            # Wait a short time and check if the file has changed
            time.sleep(2)
            
            # Check if the file still exists
            if not file_path.exists():
                return False
                
            # Get the new file size and modification time
            new_size = file_path.stat().st_size
            new_mtime = file_path.stat().st_mtime
            
            # If the file size or modification time has changed, it's still being modified
            if current_size != new_size or current_mtime != new_mtime:
                print(f"[DEBUG] File {file_path.name} is still being modified: Size changed from {current_size} to {new_size}, mtime from {current_mtime} to {new_mtime}")
                return False
                
            # Check if the file was modified in the last 45 seconds (increased for network drives)
            if time.time() - new_mtime < 45:
                print(f"[DEBUG] File {file_path.name} was modified in the last 45 seconds, waiting for stability")
                return False
                
            # Additional check: try to open the file exclusively to ensure it's not locked
            try:
                with open(file_path, 'rb') as f:
                    # Try to read a small amount to verify file is accessible
                    f.read(1024)
            except (PermissionError, IOError) as e:
                print(f"[DEBUG] File {file_path.name} is still locked or inaccessible: {str(e)}")
                return False
                
            return True
            
        except Exception as e:
            print(f"[DEBUG] Error checking if file is ready: {str(e)}")
            return False
            
    def is_folder_ready_for_processing(self, folder_path):
        """Check if a folder is ready for processing by verifying no files are being modified.
        
        Args:
            folder_path: Path to the folder to check
            
        Returns:
            bool: True if the folder is ready for processing, False otherwise
        """
        try:
            folder_path = Path(folder_path)
            if not folder_path.exists() or not folder_path.is_dir():
                return False
                
            # Check all files in the folder
            for file_path in folder_path.glob('**/*'):
                if file_path.is_file():
                    if not self.is_file_ready_for_processing(file_path):
                        return False
                        
            return True
            
        except Exception as e:
            print(f"[DEBUG] Error checking if folder is ready: {str(e)}")
            return False
            
    def are_files_ready_for_processing(self, base_folder, filenames):
        """Check if all referenced files are ready for processing.
        
        Args:
            base_folder: Base folder where files should be located
            filenames: List of filenames to check
            
        Returns:
            bool: True if all files are ready, False otherwise
        """
        try:
            base_folder = Path(base_folder)
            
            for filename in filenames:
                # Try to find the file with exact, case-insensitive, or fuzzy matching
                file_path = self.find_file_with_smart_matching(base_folder, filename)
                
                if file_path and not self.is_file_ready_for_processing(file_path):
                    print(f"[DEBUG] Referenced file {filename} is not ready for processing")
                    return False
                    
            return True
            
        except Exception as e:
            print(f"[DEBUG] Error checking if referenced files are ready: {str(e)}")
            return False
            
    def find_file_with_smart_matching(self, folder, filename):
        """Find a file using smart matching (exact, case-insensitive, fuzzy).
        
        Args:
            folder: Folder to search in
            filename: Filename to find
            
        Returns:
            Path or None: Path to the file if found, None otherwise
        """
        try:
            folder = Path(folder)
            
            # Try exact match
            file_path = folder / filename
            if file_path.exists() and file_path.is_file():
                return file_path
                
            # Try case-insensitive match
            for item in folder.glob('*'):
                if item.is_file() and item.name.lower() == filename.lower():
                    return item
                    
            # Try fuzzy matching
            best_match = None
            best_ratio = 0
            for item in folder.glob('*'):
                if item.is_file():
                    ratio = difflib.SequenceMatcher(None, filename.lower(), item.name.lower()).ratio()
                    if ratio > 0.8 and ratio > best_ratio:  # 80% similarity threshold
                        best_match = item
                        best_ratio = ratio
                        
            return best_match
            
        except Exception as e:
            print(f"[DEBUG] Error finding file with smart matching: {str(e)}")
            return None

    def browse_output_folder(self):
        """Open a folder dialog to select the output folder."""
        try:
            from tkinter import filedialog
            
            # Get the current output folder
            current_folder = self.properties.get('outbox_folder', {}).get('default', '')
            
            # Open folder dialog
            folder = filedialog.askdirectory(
                title="Select Output Folder",
                initialdir=current_folder if current_folder else None
            )
            
            if folder:
                # Update the property
                self.properties['outbox_folder']['default'] = folder
                
                # Update the display
                self.output_path_var.set(folder)
                
                # Refresh the file browser
                self.refresh_file_browser()
                
                self.update_log(f"Output folder set to: {folder}")
                
        except Exception as e:
            print(f"[DEBUG] Error browsing output folder: {str(e)}")
            self.update_log(f"Error browsing output folder: {str(e)}")

    def refresh_file_browser(self):
        """Refresh the file browser with the current output folder."""
        try:
            # Clear the treeview
            for item in self.file_tree.get_children():
                self.file_tree.delete(item)
                
            # Get the output folder
            output_folder = self.properties.get('outbox_folder', {}).get('default', '')
            if not output_folder:
                self.update_log("No output folder configured. Please set the output folder.")
                return
                
            output_folder = Path(output_folder)
            if not output_folder.exists():
                self.update_log(f"Output folder does not exist: {output_folder}")
                return
                
            # Update the path display
            self.output_path_var.set(str(output_folder))
                
            # Populate the tree
            self.populate_file_tree(output_folder, "")
            
            self.update_log(f"Refreshed file browser for: {output_folder}")
            
        except Exception as e:
            print(f"[DEBUG] Error refreshing file browser: {str(e)}")
            self.update_log(f"Error refreshing file browser: {str(e)}")

    def populate_file_tree(self, folder, parent):
        """Populate the file tree with the contents of the folder.
        
        Args:
            folder: Path to the folder to populate
            parent: Parent node in the tree
        """
        try:
            # Sort items: folders first, then files
            items = sorted(folder.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))
            
            for item in items:
                # Format size
                if item.is_file():
                    size = item.stat().st_size
                    if size < 1024:
                        size_str = f"{size} B"
                    elif size < 1024 * 1024:
                        size_str = f"{size / 1024:.1f} KB"
                    else:
                        size_str = f"{size / (1024 * 1024):.1f} MB"
                else:
                    size_str = ""
                    
                # Format modified time
                modified = datetime.fromtimestamp(item.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                
                # Add to tree
                if item.is_dir():
                    # Use a folder icon
                    folder_id = self.file_tree.insert(parent, "end", text=item.name, values=(size_str, modified), open=False, tags=("folder",))
                    # Recursively populate subfolders
                    self.populate_file_tree(item, folder_id)
                else:
                    # Use a file icon based on extension
                    if item.suffix.lower() in ('.txt', '.md'):
                        tags = ("text_file",)
                    elif item.suffix.lower() in ('.docx', '.doc'):
                        tags = ("word_file",)
                    elif item.suffix.lower() in ('.pdf'):
                        tags = ("pdf_file",)
                    elif item.suffix.lower() in ('.jpg', '.jpeg', '.png', '.gif', '.bmp'):
                        tags = ("image_file",)
                    elif item.suffix.lower() in ('.mp3', '.wav', '.m4a', '.ogg'):
                        tags = ("audio_file",)
                    elif item.suffix.lower() in ('.mp4', '.avi', '.mov', '.mkv'):
                        tags = ("video_file",)
                    else:
                        tags = ("file",)
                        
                    self.file_tree.insert(parent, "end", text=item.name, values=(size_str, modified), tags=tags)
                    
        except Exception as e:
            print(f"[DEBUG] Error populating file tree: {str(e)}")
            self.update_log(f"Error populating file tree: {str(e)}")

    def get_full_path(self, item_id):
        """Get the full path of an item in the file tree.
        
        Args:
            item_id: ID of the item in the tree
            
        Returns:
            Path: Full path to the item
        """
        try:
            # Start with the item name
            path_parts = [self.file_tree.item(item_id, "text")]
            
            # Add parent names
            parent_id = self.file_tree.parent(item_id)
            while parent_id:
                path_parts.insert(0, self.file_tree.item(parent_id, "text"))
                parent_id = self.file_tree.parent(parent_id)
                
            # Add the output folder
            output_folder = self.properties.get('outbox_folder', {}).get('default', '')
            
            # Combine all parts
            return Path(output_folder) / Path(*path_parts)
            
        except Exception as e:
            print(f"[DEBUG] Error getting full path: {str(e)}")
            self.update_log(f"Error getting full path: {str(e)}")
            return None

    def open_selected_file(self):
        """Open the selected file with the default application."""
        try:
            # Get the selected item
            selected = self.file_tree.selection()
            if not selected:
                self.update_log("No file selected.")
                return
                
            # Get the full path
            item_path = self.get_full_path(selected[0])
            
            if not item_path or not item_path.exists():
                self.update_log(f"File not found: {item_path}")
                return
                
            # Open the file with the default application
            if item_path.is_file():
                self.update_log(f"Opening file: {item_path}")
                os.startfile(str(item_path))
            else:
                # If it's a directory, expand/collapse it
                if self.file_tree.item(selected[0], "open"):
                    self.file_tree.item(selected[0], open=False)
                else:
                    self.file_tree.item(selected[0], open=True)
                
        except Exception as e:
            print(f"[DEBUG] Error opening file: {str(e)}")
            self.update_log(f"Error opening file: {str(e)}")

    def open_containing_folder(self):
        """Open the folder containing the selected file."""
        try:
            # Get the selected item
            selected = self.file_tree.selection()
            if not selected:
                self.update_log("No file selected.")
                return
                
            # Get the full path
            item_path = self.get_full_path(selected[0])
            
            if not item_path or not item_path.exists():
                self.update_log(f"Path not found: {item_path}")
                return
                
            # If it's a file, get its parent folder
            if item_path.is_file():
                folder_path = item_path.parent
            else:
                folder_path = item_path
                
            # Open the folder in Explorer
            self.update_log(f"Opening folder: {folder_path}")
            os.startfile(str(folder_path))
                
        except Exception as e:
            print(f"[DEBUG] Error opening folder: {str(e)}")
            self.update_log(f"Error opening folder: {str(e)}")

    def copy_file_path(self):
        """Copy the path of the selected file to the clipboard."""
        try:
            # Get the selected item
            selected = self.file_tree.selection()
            if not selected:
                self.update_log("No file selected.")
                return
                
            # Get the full path
            item_path = self.get_full_path(selected[0])
            
            if not item_path:
                self.update_log("Could not determine file path.")
                return
                
            # Copy to clipboard
            self.monitor_window.clipboard_clear()
            self.monitor_window.clipboard_append(str(item_path))
            
            self.update_log(f"Copied path to clipboard: {item_path}")
                
        except Exception as e:
            print(f"[DEBUG] Error copying path: {str(e)}")
            self.update_log(f"Error copying path: {str(e)}")

    def on_file_double_click(self, event):
        """Handle double-click event on a file."""
        try:
            # Get the selected item
            selected = self.file_tree.selection()
            if not selected:
                self.update_log("No file selected.")
                return
                
            # Get the full path
            item_path = self.get_full_path(selected[0])
            
            if not item_path or not item_path.exists():
                self.update_log(f"File not found: {item_path}")
                return
                
            # Open the file with the default application
            if item_path.is_file():
                self.update_log(f"Opening file: {item_path}")
                os.startfile(str(item_path))
                
        except Exception as e:
            print(f"[DEBUG] Error on file double-click: {str(e)}")
            self.update_log(f"Error on file double-click: {str(e)}")

    def on_file_right_click(self, event):
        """Handle right-click event on a file."""
        try:
            # Select the item under the cursor
            item_id = self.file_tree.identify_row(event.y)
            if item_id:
                self.file_tree.selection_set(item_id)
                self.file_context_menu.post(event.x_root, event.y_root)
                
        except Exception as e:
            print(f"[DEBUG] Error on file right-click: {str(e)}")
            self.update_log(f"Error on file right-click: {str(e)}")

    def process_video_file(self, file_path):
        """Process a video file by extracting audio and then transcribing it."""
        try:
            file_path = Path(file_path)
            project_folder = file_path.parent
            
            # Create extracted audio file path
            audio_file = project_folder / f"{file_path.stem}_extracted_audio.mp3"
            self.update_log(f"Extracting audio from video: {file_path.name}")
            print(f"[DEBUG] Extracting audio to: {audio_file}")
            
            # Extract audio from video
            if not self.extract_audio_from_video(file_path, audio_file):
                raise Exception(f"Failed to extract audio from video file: {file_path}")
            
            self.update_log(f"Audio extraction complete, processing audio: {audio_file.name}")
            
            try:
                # Process the extracted audio file
                result = self.process_audio_file(audio_file)
                
                # Delete the extracted audio file after processing
                if audio_file.exists():
                    audio_file.unlink()
                    print(f"[DEBUG] Deleted extracted audio file: {audio_file}")
                
                # Delete the original video file after processing
                if file_path.exists():
                    file_path.unlink()
                    print(f"[DEBUG] Deleted original video file: {file_path}")
                
                return result
                
            except Exception as e:
                self.update_log(f"Error processing extracted audio: {str(e)}")
                print(f"[DEBUG] Error processing extracted audio: {str(e)}")
                traceback.print_exc()
                
                # Clean up extracted audio file if there was an error
                if audio_file.exists():
                    audio_file.unlink()
                    print(f"[DEBUG] Deleted extracted audio file after error: {audio_file}")
                
                raise
        
        except Exception as e:
            self.update_log(f"Error in process_video_file: {str(e)}")
            print(f"[DEBUG] Error in process_video_file: {str(e)}")
            traceback.print_exc()
            raise
